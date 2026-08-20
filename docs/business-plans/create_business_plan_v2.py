from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Bonds brand colors
BONDS_GOLD = RGBColor(212, 168, 83)
BONDS_DARK = RGBColor(10, 15, 26)
BONDS_GRAY = RGBColor(102, 102, 102)
BONDS_BLUE = RGBColor(0, 51, 102)

def set_rtl(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.right_to_left = True

def add_heading_rtl(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    set_rtl(heading)
    for run in heading.runs:
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        if level == 1:
            run.font.color.rgb = BONDS_BLUE
            run.font.size = Pt(18)
        elif level == 2:
            run.font.color.rgb = BONDS_GOLD
            run.font.size = Pt(14)
    return heading

def add_paragraph_rtl(doc, text, bold=False, italic=False, size=12, color=RGBColor(0, 0, 0)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    set_rtl(p)
    return p

def add_bullet_rtl(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(11)
    set_rtl(p)
    return p

def add_table_rtl(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table.autofit = True
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            set_rtl(paragraph)
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                run.font.color.rgb = RGBColor(255, 255, 255)
        # Set header cell background
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '0A0F1A')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                set_rtl(paragraph)
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    return table

def add_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.right_to_left = True
    
    run1 = p.add_run('Bonds Financial Management & Consulting  |  ')
    run1.font.name = 'Arial'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run1.font.size = Pt(9)
    run1.font.color.rgb = BONDS_GRAY
    
    run2 = p.add_run('www.bonds-global.com')
    run2.font.name = 'Arial'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run2.font.size = Pt(9)
    run2.font.color.rgb = BONDS_GOLD

# Create document
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

# Section setup for footer
section = doc.sections[0]
section.page_height = Inches(11.69)
section.page_width = Inches(8.27)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
add_page_number(section)

# ===== COVER PAGE =====
# Header on cover
header_p = doc.add_paragraph()
header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
header_p.paragraph_format.right_to_left = True
run = header_p.add_run('Bonds Financial Management & Consulting')
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = BONDS_GOLD

subheader = doc.add_paragraph()
subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
subheader.paragraph_format.right_to_left = True
run = subheader.add_run('Bonds Global')
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
run.font.size = Pt(12)
run.font.color.rgb = BONDS_DARK

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Main title box
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.right_to_left = True
run = title_p.add_run('خطة عمل')
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = BONDS_BLUE

title2_p = doc.add_paragraph()
title2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title2_p.paragraph_format.right_to_left = True
run = title2_p.add_run('شركة إحياء الأصول غير المكتملة العقارية')
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = BONDS_GOLD

doc.add_paragraph()
doc.add_paragraph()

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.right_to_left = True
run = subtitle.add_run('شركة مساهمة مقفلة')
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
run.font.size = Pt(16)
run.font.bold = True
run.font.color.rgb = BONDS_DARK

subtitle2 = doc.add_paragraph()
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle2.paragraph_format.right_to_left = True
run = subtitle2.add_run('المملكة العربية السعودية')
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
run.font.size = Pt(14)
run.font.color.rgb = BONDS_GRAY

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Date
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.paragraph_format.right_to_left = True
run = date_p.add_run('يونيو 2026')
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
run.font.size = Pt(12)
run.font.color.rgb = BONDS_GRAY

conf_p = doc.add_paragraph()
conf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
conf_p.paragraph_format.right_to_left = True
run = conf_p.add_run('وثيقة سرية - مخصصة للمساهمين المؤسسين')
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = BONDS_GRAY

doc.add_page_break()

# ===== EXECUTIVE SUMMARY =====
add_heading_rtl(doc, 'ملخص تنفيذي', 1)
add_paragraph_rtl(doc, 'تأسست شركة إحياء الأصول غير المكتملة العقارية كنموذج استثماري متخصص، يهدف إلى إحياء العقارات غير المكتملة في المملكة العربية السعودية من خلال شراكات محاصة مع مالكيها، دون الحاجة إلى شراء الأراضي أو الاعتماد على التمويل الخارجي.')
add_paragraph_rtl(doc, 'تعمل الشركة بصورة قانونية كنظام مساهمة مقفلة، ممولة بالكامل من المساهمين المؤسسين، وتلتزم بأعلى معايير الكود السعودي للبناء والتشطيب، مع تغطية تأمينية شاملة لجميع المشاريع.')

add_heading_rtl(doc, 'أبرز المؤشرات المالية', 2)
rows = [
    ['رأس المال المصرح', '200 مليون ريال'],
    ['إجمالي الاستثمار المستهدف', '226 مليون ريال'],
    ['الاستثمار الفعلي المبدئي', '25 مليون ريال'],
    ['العائد الداخلي المتوقع (IRR) على أفق 10 سنوات', '19% - 20%'],
    ['مضاعفة رأس المال عند الخروج', '2.5x - 4x'],
    ['فترة استرداد رأس المال المستدعى', '5.5 سنة'],
]
add_table_rtl(doc, ['المؤشر', 'القيمة'], rows)

add_heading_rtl(doc, 'الفرصة السوقية', 2)
add_bullet_rtl(doc, 'السعودية تشهد نمواً عقارياً غير مسبوق بفضل رؤية 2030 ومشاريعها التنموية الضخمة.')
add_bullet_rtl(doc, 'تتوفر في السوق آلاف العقارات غير المكتملة نتيجة نقص السيولة أو الخبرة التنفيذية.')
add_bullet_rtl(doc, 'الطلب على الوحدات الجاهزة يفوق العرض في المدن الرئيسية.')
add_bullet_rtl(doc, 'تسجل الرياض وجدة عوائد إيجارية مرتفعة تصل إلى 8.89% و 7.89% على التوالي.')

add_heading_rtl(doc, 'نموذج العمل', 2)
add_bullet_rtl(doc, 'الشراكة بالمحاصة مع مالك العقار غير المكتمل.')
add_bullet_rtl(doc, 'تتحمل الشركة تكاليف التشطيب والتسويق والبيع.')
add_bullet_rtl(doc, 'يحتفظ المالك بملكية عقاره طوال فترة المشروع.')
add_bullet_rtl(doc, 'يتم توزيع الأرباح بين الشركة والمالك وفقاً للعقد الموقع.')

doc.add_page_break()

# 1. Company Overview
add_heading_rtl(doc, '1. نظرة عامة على الشركة', 1)
add_heading_rtl(doc, '1.1 الرؤية', 2)
add_paragraph_rtl(doc, 'أن نكون الشريك الاستراتيجي المفضل لمالكي العقارات غير المكتملة في المملكة العربية السعودية، من خلال تقديم حلول احترافية لإحياء أصولهم وتحقيق أقصى قيمة ممكنة لهم وللمساهمين.')

add_heading_rtl(doc, '1.2 الرسالة', 2)
add_paragraph_rtl(doc, 'إحياء العقارات غير المكتملة بمعايير عالية من الجودة والالتزام، وبنموذج مالي مستدام يحقق العوائد للشركاء ويحفظ حقوق ملاك الأصول.')

add_heading_rtl(doc, '1.3 القيم الأساسية', 2)
add_bullet_rtl(doc, 'الاحترام: احترام مالك الأصل وحقوقه الملكية في كل مرحلة من مراحل المشروع.')
add_bullet_rtl(doc, 'الجودة: الالتزام بالكود السعودي وأعلى معايير التشطيب والسلامة.')
add_bullet_rtl(doc, 'الشفافية: وضوح كامل في العقود والتقارير المالية والتشغيلية.')
add_bullet_rtl(doc, 'الاستدامة: بناء نموذج اقتصادي طويل الأجل يحقق القيمة لجميع الأطراف.')
add_bullet_rtl(doc, 'الابتكار: استخدام التقنية والحلول الذكية لتحسين الكفاءة وتقليل المخاطر.')

add_heading_rtl(doc, '1.4 الشكل القانوني', 2)
add_paragraph_rtl(doc, 'تأسست الشركة بصورة قانونية كنظام مساهمة مقفلة، مرخصة من وزارة التجارة وهيئة العقار (REGA)، وفقاً لنظام الشركات السعودي. يتيح هذا الشكل جمع رأس المال من المساهمين المؤسسين مع الحفاظ على مرونة الإدارة وسرعة اتخاذ القرار.')

add_heading_rtl(doc, '1.5 رأس المال وهيكل الملكية', 2)
rows = [
    ['رأس المال المصرح', '200,000,000 ريال سعودي'],
    ['عدد الأسهم', '40,000 سهم'],
    ['القيمة الاسمية للسهم', '5,000 ريال سعودي'],
    ['القيمة الفعلية للسهم', '5,650 ريال سعودي'],
    ['علاوة الإصدار ومصاريف التأسيس', '26,000,000 ريال سعودي'],
    ['إجمالي الاستثمار المستهدف', '226,000,000 ريال سعودي'],
]
add_table_rtl(doc, ['البند', 'القيمة'], rows)

doc.add_page_break()

# 2. Market Analysis
add_heading_rtl(doc, '2. تحليل السوق', 1)
add_heading_rtl(doc, '2.1 نظرة عامة على السوق العقاري السعودي', 2)
add_paragraph_rtl(doc, 'يشهد قطاع الاستثمار والتطوير العقاري في المملكة العربية السعودية تحولاً جذرياً، مدفوعاً برؤية 2030 والمشاريع التنموية الضخمة. يتجه القطاع نحو النضج، مع بدء ديناميكيات طلب حقيقية وملموسة في سلوك المستهلكين والمستثمرين.')

add_heading_rtl(doc, '2.2 المؤشرات الرئيسية', 2)
rows = [
    ['متوسط العائد الإيجاري في الرياض', '8.89%'],
    ['متوسط العائد الإيجاري في جدة', '7.89%'],
    ['معدل نمو الصفقات العقارية (2025)', '+14.2%'],
    ['حجم التمويل العقاري الإجمالي', '730 مليار ريال'],
    ['متوسط فجوة السعر الفعلي مقارنة بالطلب', '6%'],
    ['علاوة البناء الجديد', '+25%'],
]
add_table_rtl(doc, ['المؤشر', 'القيمة'], rows)

add_heading_rtl(doc, '2.3 محركات النمو', 2)
add_bullet_rtl(doc, 'رؤية 2030 والمشاريع العملاقة مثل نيوم، الدرعية، البحر الأحمر، وروشن.')
add_bullet_rtl(doc, 'برامج التمويل السكني "سكني" الموجهة للوحدات المتوسطة (70 - 120 م²).')
add_bullet_rtl(doc, 'نمو الرواتب وتوجه الأفراد نحو المساحات الأصغر والأكثر كفاءة.')
add_bullet_rtl(doc, 'الهجرة الداخلية نحو المدن الكبرى والعاصمة الرياض.')
add_bullet_rtl(doc, 'انتعاش سوق الإيجارات وارتفاع عوائدها إلى مستويات جاذبة.')

add_heading_rtl(doc, '2.4 المناطق المستهدفة', 2)
add_paragraph_rtl(doc, 'مدينة الرياض:', bold=True)
add_bullet_rtl(doc, 'الشمال الفاخر: العقيق، حطين، الملقا - يجذب التنفيذيين ويوفر عوائد إيجارية مرتفعة.')
add_bullet_rtl(doc, 'الشمال الشرقي: المونسية، الرمال، الياسمين، النرجس - كثافة سكانية عالية وبنية تحتية حديثة.')

add_paragraph_rtl(doc, 'مدينة جدة:', bold=True)
add_bullet_rtl(doc, 'الساحل الشمالي: أبحر الشمالية، الصواري - يمثل مستقبلاً سكنياً وسياحياً واعداً.')
add_bullet_rtl(doc, 'الوسط الحيوي: الحمراء، النعيم، الزهراء، السلامة - يوفر عوائد إيجارية ثابتة ونشاطاً تجارياً مستمراً.')

add_heading_rtl(doc, '2.5 تحليل المنافسة', 2)
add_paragraph_rtl(doc, 'المنافسة الرئيسية ليست من الشركات الكبرى، بل من المطورين التقليديين الذين يتطلب نموذجهم شراء أراضٍ ورأس مال ضخم. تكمن ميزتنا التنافسية في عدم الحاجة إلى شراء الأرض، والتركيز على الأصول غير المكتملة التي لا تجذب المطورين الكبار.')

doc.add_page_break()

# 3. Services
add_heading_rtl(doc, '3. الخدمات والحلول', 1)
add_heading_rtl(doc, '3.1 الخدمات الأساسية', 2)
add_bullet_rtl(doc, 'إجراء تقييم هندسي وقانوني شامل للعقار غير المكتمل.')
add_bullet_rtl(doc, 'إعداد التصاميم المعمارية والإنشائية والميكانيكية المطلوبة.')
add_bullet_rtl(doc, 'إدارة عمليات التشطيب والإشراف اليومي على التنفيذ.')
add_bullet_rtl(doc, 'التسويق والبيع الاحترافي للوحدات بأسعار السوق.')
add_bullet_rtl(doc, 'إدارة عملية التسليم ونقل الملكية وتوزيع الأرباح.')

add_heading_rtl(doc, '3.2 نماذج الشراكة', 2)
rows = [
    ['النموذج', 'الوصف', 'الحالة المناسبة'],
    ['المحاصة الكاملة', 'تتولى الشركة التشطيب والتسويق مقابل نسبة متفق عليها من الأرباح', 'المالك يملك سيولة محدودة ويرغب في إحياء العقار'],
    ['الشراء الجزئي', 'تشتري الشركة جزءاً من الوحدات بعد إتمام التشطيب', 'المالك يرغب في تحقيق سيولة فورية'],
    ['رسوم الإدارة', 'تتقاضى الشركة رسوم إدارة ثابتة مقابل التشطيب والتسويق', 'المالك يرغب في الاحتفاز بمعظم الأرباح'],
    ['النموذج المختلط', 'مزيج متوازن من المحاصة والشراء ورسوم الإدارة', 'معظم الحالات العملية'],
]
add_table_rtl(doc, ['النموذج', 'الوصف', 'الحالة المناسبة'], rows)

doc.add_page_break()

# 4. Business Model
add_heading_rtl(doc, '4. النموذج التجاري', 1)
add_heading_rtl(doc, '4.1 آلية عمل الشركة', 2)
add_bullet_rtl(doc, 'البحث عن عقارات غير مكتملة في المواقع المستهدفة وتحديد فرصها.')
add_bullet_rtl(doc, 'إجراء التقييم الهندسي والقانوني والمالي الشامل.')
add_bullet_rtl(doc, 'التفاوض مع المالك وإبرام عقد محاصة أو شراكة واضح.')
add_bullet_rtl(doc, 'إنجاز أعمال التشطيب وفق أعلى المعايير والكود السعودي.')
add_bullet_rtl(doc, 'تسويق الوحدات وبيعها بأسعار السوق المناسبة.')
add_bullet_rtl(doc, 'توزيع الأرباح بين الشركة والمالك وفق بنود العقد.')

add_heading_rtl(doc, '4.2 مصادر الإيرادات', 2)
rows = [
    ['مصدر الإيراد', 'النسبة المتوقعة'],
    ['إيرادات المحاصة (نصيب الشركة)', '60%'],
    ['أرباح إعادة البيع والشراء', '25%'],
    ['رسوم الإدارة والتسويق', '10%'],
    ['خدمات استشارية وإدارية', '5%'],
]
add_table_rtl(doc, ['مصدر الإيراد', 'النسبة المتوقعة'], rows)

add_heading_rtl(doc, '4.3 هيكل التكاليف', 2)
rows = [
    ['البند', 'النسبة من تكلفة المشروع'],
    ['تكاليف التشطيب المباشرة', '55%'],
    ['التراخيص والرسوم البلدية', '4%'],
    ['الاستشارات الهندسية والإشراف', '5%'],
    ['التسويق والمبيعات', '7%'],
    ['احتياطي المخاطر', '10%'],
    ['التكاليف الإدارية التشغيلية', '6%'],
    ['رسوم البيع والنقل', '9%'],
    ['التأمين', '4%'],
]
add_table_rtl(doc, ['البند', 'النسبة'], rows)

add_heading_rtl(doc, '4.4 نقاط القوة التنافسية', 2)
add_bullet_rtl(doc, 'لا يتطلب النموذج شراء أراضٍ، مما يقلل رأس المال المطلوب للبدء.')
add_bullet_rtl(doc, 'سرعة الدخول للسوق مقارنة بالتطوير العقاري التقليدي.')
add_bullet_rtl(doc, 'شبكة علاقات واسعة مع مالكين، مصفين، محامين، ووسطاء عقاريين.')
add_bullet_rtl(doc, 'الالتزام بالكود السعودي والتأمين، مما يقلل المخاطر القانونية والمالية.')
add_bullet_rtl(doc, 'نموذج مالي مرن وممول ذاتياً من المساهمين المؤسسين.')

doc.add_page_break()

# 5. Marketing & Sales
add_heading_rtl(doc, '5. استراتيجية التسويق والمبيعات', 1)
add_heading_rtl(doc, '5.1 استراتيجية التسويق الرقمي', 2)
rows = [
    ['القناة', 'الهدف', 'الميزانية الشهرية'],
    ['Google Ads', 'الوصول إلى مالكي العقارات غير المكتملة', '10,000 ريال'],
    ['Snapchat Ads', 'استهداف الفئة العمرية 25 - 45 سنة', '8,000 ريال'],
    ['Instagram / TikTok', 'عرض مشاريع "قبل وبعد"', '5,000 ريال'],
    ['SEO', 'الظهور في نتائج البحث عن إحياء المباني', '3,000 ريال'],
    ['LinkedIn', 'الوصول إلى المستثمرين ومالكي العقارات', '2,000 ريال'],
]
add_table_rtl(doc, ['القناة', 'الهدف', 'الميزانية الشهرية'], rows)

add_heading_rtl(doc, '5.2 قمع المبيعات', 2)
add_bullet_rtl(doc, 'الوعي: الوصول إلى 100,000 شخص شهرياً.')
add_bullet_rtl(doc, 'الاهتمام: تحويل 5,000 شخص إلى مهتمين بالخدمة.')
add_bullet_rtl(doc, 'التقييم: التواصل مع 500 مالك محتمل شهرياً.')
add_bullet_rtl(doc, 'النية: عقد 50 اجتماعاً مبدئياً شهرياً.')
add_bullet_rtl(doc, 'الشراء: إبرام 2 - 4 عقود محاصة شهرياً.')

add_heading_rtl(doc, '5.3 استراتيجية الوصول إلى المالكين', 2)
add_bullet_rtl(doc, 'بناء قاعدة بيانات شاملة للعقارات غير المكتملة في المناطق المستهدفة.')
add_bullet_rtl(doc, 'إبرام شراكات استراتيجية مع المصفين والوسطاء العقاريين.')
add_bullet_rtl(doc, 'المشاركة في المعارض العقارية السنوية داخل المملكة.')
add_bullet_rtl(doc, 'إطلاق حملات توعوية حول فوائد إحياء الأصول غير المكتملة.')

doc.add_page_break()

# 6. Operational Plan
add_heading_rtl(doc, '6. الخطة التشغيلية', 1)
add_heading_rtl(doc, '6.1 مراحل المشروع الواحد', 2)
rows = [
    ['المرحلة', 'المدة', 'الأنشطة الرئيسية'],
    ['التقييم', '4 أسابيع', 'فحص هيكلي، مراجعة قانونية، تقدير تكلفة مبدئي'],
    ['التعاقد', '2 - 4 أسابيع', 'صياغة عقد المحاصة، استصدار الموافقات اللازمة'],
    ['التصميم', '4 - 8 أسابيع', 'إعداد التصاميم واستصدار التراخيص البلدية'],
    ['التشطيب', '16 - 24 أسبوع', 'اختيار مقاول، إشراف يومي، تفتيش دوري'],
    ['التسويق والبيع', '12 - 24 أسبوع', 'عرض الوحدات، استقبال العملاء، عقود البيع'],
    ['التسليم والتسوية', '4 - 8 أسابيع', 'تسليم الوحدات، نقل الملكية، توزيع الأرباح'],
]
add_table_rtl(doc, ['المرحلة', 'المدة', 'الأنشطة الرئيسية'], rows)

add_heading_rtl(doc, '6.2 مؤشرات الأداء الرئيسية (KPIs)', 2)
rows = [
    ['المؤشر', 'الهدف', 'التحذير', 'الخطر'],
    ['نسبة إنجاز التشطيب', 'أكثر من 80% في الشهر الرابع', 'أقل من 60%', 'تأخر المقاول'],
    ['معدل بيع الوحدات', 'أكثر من 30% في الشهر الثاني', 'أقل من 10%', 'ضعف الطلب'],
    ['تكلفة التشطيب مقارنة بالمتوقع', 'أقل من 105%', 'أكثر من 115%', 'عيوب مخفية'],
    ['رضا العملاء (NPS)', 'أكثر من 50', 'أقل من 30', 'مشاكل في الجودة'],
    ['مدة البيع', 'أقل من 6 أشهر', 'أكثر من 9 أشهر', 'ركود في السوق'],
]
add_table_rtl(doc, ['المؤشر', 'الهدف', 'التحذير', 'الخطر'], rows)

add_heading_rtl(doc, '6.3 نظام الجودة', 2)
add_bullet_rtl(doc, 'الالتزام الكامل بكود البناء السعودي في جميع مراحل التنفيذ.')
add_bullet_rtl(doc, 'إجراء فحوصات جودة دورية في كل مرحلة من مراحل التشطيب.')
add_bullet_rtl(doc, 'توثيق جميع المواد المستخدمة وشهادات المطابقة الخاصة بها.')
add_bullet_rtl(doc, 'تسليم المشروع بملف فني وقانوني كامل للمالك والمشتري.')

doc.add_page_break()

# 7. Management Team
add_heading_rtl(doc, '7. الفريق الإداري والتنظيمي', 1)
add_heading_rtl(doc, '7.1 الهيكل التنظيمي', 2)
add_paragraph_rtl(doc, 'تم تصميم الهيكل الإداري ليكون مرناً وفعالاً، مع بدء تشغيلي يعتمد على 15 موظفاً في السنوات الأولى، مع إمكانية التوسع تبعاً لنمو المحفظة.')

add_heading_rtl(doc, '7.2 الأدوار الإدارية الرئيسية', 2)
rows = [
    ['الدور', 'المسؤوليات الرئيسية'],
    ['الرئيس التنفيذي', 'وضع الاستراتيجية العامة، إدارة علاقات المستثمرين، اتخاذ القرارات الكبرى'],
    ['المدير التشغيلي', 'إدارة المشاريع، التنسيق مع المقاولين، ضمان جودة التنفيذ'],
    ['المدير المالي', 'إدارة التدفقات النقدية، إعداد التقارير المالية، التخطيط المالي'],
    ['المدير التجاري', 'إدارة التسويق والمبيعات، بناء علاقات المالكين، تطوير الأعمال'],
    ['المدير القانوني', 'صياغة العقود، ضمان الامتثال، حل النزاعات'],
    ['مدير الموارد البشرية والإدارة', 'إدارة الموظفين والإجراءات الإدارية'],
]
add_table_rtl(doc, ['الدور', 'المسؤوليات الرئيسية'], rows)

add_heading_rtl(doc, '7.3 خطة التوظيف', 2)
rows = [
    ['السنة', 'عدد الموظفين', 'الملاحظات'],
    ['السنة الأولى', '10 - 15', 'فريق أساسي لتشغيل المرحلة الأولى'],
    ['السنة الثانية', '15', 'تثبيت الهيكل التشغيلي'],
    ['السنة الثالثة', '18 - 20', 'التوسع في عدد المشاريع'],
    ['السنة الخامسة', '25 - 30', 'دعم محفظة 10 مشاريع سنوياً'],
]
add_table_rtl(doc, ['السنة', 'عدد الموظفين', 'الملاحظات'], rows)

doc.add_page_break()

# 8. Financial Plan
add_heading_rtl(doc, '8. الخطة المالية', 1)
add_heading_rtl(doc, '8.1 افتراضات النموذج المالي', 2)
rows = [
    ['الافتراض', 'القيمة'],
    ['متوسط مساحة المبنى', '1,500 متر مربع'],
    ['عدد الوحدات في المبنى', '6 - 8 وحدات'],
    ['تكلفة التشطيب', '1,200 - 1,800 ريال لكل متر مربع'],
    ['سعر البيع', '2,500 - 3,500 ريال لكل متر مربع'],
    ['نسبة المحاصة للشركة', '50%'],
]
add_table_rtl(doc, ['الافتراض', 'القيمة'], rows)

add_heading_rtl(doc, '8.2 استخدامات رأس المال', 2)
rows = [
    ['الاستخدام', 'المبلغ', 'النسبة'],
    ['استثمارات في مشاريع (شراء + تشطيب)', '150 مليون ريال', '66%'],
    ['رأس المال العامل التشغيلي', '40 مليون ريال', '18%'],
    ['تكنولوجيا وأنظمة', '10 مليون ريال', '4%'],
    ['تأسيس وتسويق وعلاقات', '15 مليون ريال', '7%'],
    ['احتياطي سيولة طوارئ', '11 مليون ريال', '5%'],
]
add_table_rtl(doc, ['الاستخدام', 'المبلغ', 'النسبة'], rows)

add_heading_rtl(doc, '8.3 ملخص النتائج المالية المتوقعة', 2)
rows = [
    ['المؤشر', 'القيمة'],
    ['العائد الداخلي المتوقع (IRR) على أفق 10 سنوات', '19% - 20%'],
    ['مضاعفة رأس المال عند الخروج', '2.5x - 4x'],
    ['فترة استرداد رأس المال المستدعى', '5.5 سنة'],
    ['صافي الربح المتوقع في السنة الخامسة', '9.8 مليون ريال'],
    ['القيمة السوقية المتوقعة بعد 5 سنوات', '120 - 155 مليون ريال'],
]
add_table_rtl(doc, ['المؤشر', 'القيمة'], rows)

add_heading_rtl(doc, '8.4 آلية سحب رأس المال', 2)
add_paragraph_rtl(doc, 'لا يُضخ رأس المال دفعة واحدة، بل يتم سحبه على أربع دفعات مرتبطة بإنجاز مراحل محددة:')
add_bullet_rtl(doc, 'الدفعة الأولى (30%): التأسيس وشراء 3 - 5 مبانٍ أولى.')
add_bullet_rtl(doc, 'الدفعة الثانية (30%): إنجاز أول مشروعين وتحقيق أولى الإيرادات.')
add_bullet_rtl(doc, 'الدفعة الثالثة (25%): الوصول إلى 6 مشاريع نشطة وتحقيق التعادل التشغيلي.')
add_bullet_rtl(doc, 'الدفعة الرابعة (15%): احتياطي مرن يُطلق حسب الحاجة أو يُعاد للمساهمين.')

doc.add_page_break()

# 9. Risk Analysis
add_heading_rtl(doc, '9. تحليل المخاطر والتحوطات', 1)
add_heading_rtl(doc, '9.1 المخاطر والتحوطات المقترحة', 2)
rows = [
    ['المخاطر', 'الاحتمالية', 'التأثير', 'التحوط'],
    ['صعوبة الحصول على أصول جيدة', 'عالية', 'توقف النمو', 'بناء شبكة علاقات وقاعدة بيانات وشراكات مع مصفين'],
    ['عيوب إنشائية مخفية', 'عالية', 'خسارة مالية كبيرة', 'فحص هندسي شامل + تأمين + احتياطي 10%'],
    ['تأخر بيع الوحدات', 'متوسطة', 'حمل تكاليف إضافي', 'تسعير تنافسي + عقود مسبقة البيع'],
    ['تأخر المقاول في التنفيذ', 'متوسطة', 'تكاليف إضافية', 'كفالة بنكية + عقوبات تأخير في العقد'],
    ['تغير أسعار مواد البناء', 'متوسطة', 'ارتفاع التكاليف', 'عقود ثابتة السعر مع مقاولين موثوقين'],
    ['نزاعات قانونية مع المالك', 'متوسطة', 'توقف المشروع', 'عقود واضحة + إدراج شرط تحكيم ملزم'],
    ['تراجع أسعار العقار', 'منخفضة إلى متوسطة', 'انخفاض الهوامش', 'تنويع جغرافي واستهداف فئات متعددة'],
]
add_table_rtl(doc, ['المخاطر', 'الاحتمالية', 'التأثير', 'التحوط'], rows)

add_heading_rtl(doc, '9.2 سياسة التأمين', 2)
add_bullet_rtl(doc, 'تأمين المقاولين الشامل (CAR) لكل مشروع على حدة.')
add_bullet_rtl(doc, 'تأمين المسؤولية المدنية العامة للشركة.')
add_bullet_rtl(doc, 'تأمين العيوب الخفية للمباني غير المكتملة.')
add_bullet_rtl(doc, 'تأمين المسؤولية المهنية للمهندسين والاستشاريين.')
add_bullet_rtl(doc, 'تأمين مقر الشركة والأصول الثابتة.')

add_heading_rtl(doc, '9.3 الامتثال للكود السعودي', 2)
add_bullet_rtl(doc, 'كود البناء السعودي (SBC) - معايير الإنشاء والسلامة.')
add_bullet_rtl(doc, 'اشتراطات البلدية - التراخيص والتفتيش الدوري.')
add_bullet_rtl(doc, 'اشتراطات الدفاع المدني - السلامة من الحريق.')
add_bullet_rtl(doc, 'اشتراطات الهيئة العامة للعقار (REGA).')

doc.add_page_break()

# 10. Implementation Milestones
add_heading_rtl(doc, '10. خطة التنفيذ والمراحل الزمنية', 1)
add_heading_rtl(doc, '10.1 المرحلة الأولى: إثبات النموذج (السنة 1 - 2)', 2)
add_bullet_rtl(doc, 'إتمام التأسيس القانوني واستخراج التراخيص اللازمة.')
add_bullet_rtl(doc, 'ضخ 25 مليون ريال كاستثمار أولي للمرحلة التجريبية.')
add_bullet_rtl(doc, 'تنفيذ 3 - 5 مشاريع قياسية لإثبات جدوى النموذج.')
add_bullet_rtl(doc, 'بناء فريق تشغيلي أساسي يتراوح بين 10 و 15 موظفاً.')
add_bullet_rtl(doc, 'إنشاء قاعدة بيانات شاملة للأصول غير المكتملة.')

add_heading_rtl(doc, '10.2 المرحلة الثانية: النمو والتكرار (السنة 3 - 5)', 2)
add_bullet_rtl(doc, 'زيادة رأس المال المستدعى تدريجياً حسب نمو المحفظة.')
add_bullet_rtl(doc, 'الوصول إلى 6 - 10 مشاريع سنوياً.')
add_bullet_rtl(doc, 'تحقيق الربحية التشغيلية المستدامة.')
add_bullet_rtl(doc, 'تطوير نظام ERP وأتمتة العمليات التشغيلية.')
add_bullet_rtl(doc, 'التركيز على مدينتي الرياض وجدة كأسواق رئيسية.')

add_heading_rtl(doc, '10.3 المرحلة الثالثة: الخروج الاستراتيجي (السنة 5 - 7)', 2)
add_bullet_rtl(doc, 'بيع حصة مسيطرة لصندوق استثماري عقاري أو مطور كبير.')
add_bullet_rtl(doc, 'أو إجراء طرح جزئي في السوق المالية السعودية (Tadawul).')
add_bullet_rtl(doc, 'أو إعادة شراء الأسهم من المساهمين تدريجياً من الأرباح.')

add_heading_rtl(doc, '10.4 المخطط الزمني الرئيسي', 2)
rows = [
    ['المرحلة', 'المدة', 'المعالم الرئيسية'],
    ['التأسيس', '0 - 6 أشهر', 'استخراج التراخيص، تجهيز المقر، تجميع الفريق الأساسي'],
    ['الإطلاق', '6 - 12 شهراً', 'تنفيذ أول 3 مشاريع، تحقيق أولى الإيرادات'],
    ['النمو', 'السنة 2 - 3', '6 مشاريع سنوياً، الوصول للتعادل التشغيلي'],
    ['التوسع', 'السنة 4 - 5', '10 مشاريع سنوياً، أرباح مستدامة'],
    ['الخروج', 'السنة 5 - 7', 'بيع استراتيجي أو طرح جزئي'],
]
add_table_rtl(doc, ['المرحلة', 'المدة', 'المعالم الرئيسية'], rows)

doc.add_page_break()

# 11. Conclusion
add_heading_rtl(doc, '11. الخلاصة والتوصيات', 1)
add_paragraph_rtl(doc, 'تقدم شركة إحياء الأصول غير المكتملة العقارية فرصة استثمارية فريدة في السوق العقاري السعودي. يعتمد النموذج التجاري على الشراكة المحاصة مع مالكي العقارات غير المكتملة، مما يلغي الحاجة إلى شراء الأراضي ويقلل رأس المال المطلوب للبدء.')
add_paragraph_rtl(doc, 'بفضل التمويل الذاتي من المساهمين المؤسسين، والالتزام بالكود السعودي، والتغطية التأمينية الشاملة، تتمتع الشركة بمخاطر محسوبة وعوائد جذابة تصل إلى 19 - 20% كعائد داخلي متوقع على أفق 10 سنوات.')

add_heading_rtl(doc, 'التوصيات الرئيسية', 2)
add_bullet_rtl(doc, 'البدء بمرحلة تجريبية بـ 25 مليون ريال وتنفيذ 3 - 5 مشاريع قياسية.')
add_bullet_rtl(doc, 'التركيز على مدينة واحدة (الرياض أو جدة) في السنوات الأولى لتجنب تشتيت الجهد.')
add_bullet_rtl(doc, 'بناء قاعدة بيانات للأصول غير المكتملة قبل الإطلاق الرسمي.')
add_bullet_rtl(doc, 'إبرام شراكات مع مقاولين معتمدين يقدمون كفالات بنكية على أعمالهم.')
add_bullet_rtl(doc, 'إدراج شرط تحكيم ملزم في جميع عقود المحاصة.')
add_bullet_rtl(doc, 'تطوير نظام ERP لإدارة المحفظة والتدفقات النقدية.')
add_bullet_rtl(doc, 'إعداد عقد محاصة موحد يحفظ حقوق الشركة والمالك على حد سواء.')

# Save
doc.save('C:/Users/vip/bonds-global-web/docs/business-plans/business-plan-arabic-v2.docx')
print('Business plan v2 created successfully')
