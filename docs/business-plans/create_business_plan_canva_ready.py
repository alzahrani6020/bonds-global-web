from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import os
import shutil

BONDS_GOLD = RGBColor(212, 168, 83)
BONDS_DARK = RGBColor(10, 15, 26)
BONDS_GRAY = RGBColor(102, 102, 102)
BONDS_BLUE = RGBColor(0, 51, 102)
BONDS_LIGHT_GOLD = RGBColor(240, 201, 106)

AR_NAME = 'د. طلال بن حسن الزهراني'
AR_TITLE = 'المستشار المالي'
EN_NAME = 'Dr. Talal bin Hassan Al-Zahrani'
EN_TITLE = 'Financial Advisor'
PHONE = '+966 56 756 6616'
EMAIL = 'info@bonds-global.com'

def convert_logo():
    logo_webp = r'C:\Users\vip\bonds-global-web\assets\bonds-logo-2026-v2.webp'
    logo_png = r'C:\Users\vip\bonds-global-web\docs\business-plans\bonds-logo.png'
    if not os.path.exists(logo_png):
        img = Image.open(logo_webp)
        img.save(logo_png, 'PNG')
    return logo_png

def create_qr():
    import qrcode
    qr_path = r'C:\Users\vip\bonds-global-web\docs\business-plans\bonds-qr-code.png'
    if not os.path.exists(qr_path):
        qr = qrcode.QRCode(version=3, box_size=10, border=2)
        qr.add_data('https://www.bonds-global.com')
        qr.make(fit=True)
        img = qr.make_image(fill_color="#0A0F1A", back_color="white")
        img.save(qr_path)
    return qr_path

def set_rtl(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.right_to_left = True

def set_ltr(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.right_to_left = False

def add_heading_rtl(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    set_rtl(heading)
    for run in heading.runs:
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        if level == 1:
            run.font.color.rgb = BONDS_BLUE
            run.font.size = Pt(18)
        elif level == 2:
            run.font.color.rgb = BONDS_GOLD
            run.font.size = Pt(14)
    return heading

def add_heading_ltr(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    set_ltr(heading)
    for run in heading.runs:
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        if level == 1:
            run.font.color.rgb = BONDS_BLUE
            run.font.size = Pt(18)
        elif level == 2:
            run.font.color.rgb = BONDS_GOLD
            run.font.size = Pt(14)
    return heading

def add_paragraph_rtl(doc, text, bold=False, italic=False, size=12, color=RGBColor(0, 0, 0)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    set_rtl(p)
    return p

def add_paragraph_ltr(doc, text, bold=False, italic=False, size=12, color=RGBColor(0, 0, 0)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    set_ltr(p)
    return p

def add_bullet_rtl(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(11)
    set_rtl(p)
    return p

def add_bullet_ltr(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(11)
    set_ltr(p)
    return p

def add_table_rtl(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table.autofit = True
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            set_rtl(paragraph)
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                run.font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '0A0F1A')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                set_rtl(paragraph)
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    return table

def add_table_ltr(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.autofit = True
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            set_ltr(paragraph)
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                run.font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '0A0F1A')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                set_ltr(paragraph)
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    return table

def add_footer_rtl(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.right_to_left = True
    run = p.add_run(f'{AR_TITLE} | {AR_NAME} | {PHONE} | {EMAIL}')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(9)
    run.font.color.rgb = BONDS_GRAY

def add_footer_ltr(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.right_to_left = False
    run = p.add_run(f'{EN_TITLE} | {EN_NAME} | {PHONE} | {EMAIL}')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(9)
    run.font.color.rgb = BONDS_GRAY

def create_arabic_business_plan(logo_path, qr_path, output_path):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_footer_rtl(section)

    # Cover page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(logo_path, width=Inches(2.0))

    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.right_to_left = True
    run = title_p.add_run('خطة عمل')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(40)
    run.font.bold = True
    run.font.color.rgb = BONDS_BLUE

    title2_p = doc.add_paragraph()
    title2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title2_p.paragraph_format.right_to_left = True
    run = title2_p.add_run('شركة إحياء الأصول غير المكتملة العقارية')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = BONDS_GOLD

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.right_to_left = True
    run = subtitle.add_run('شركة مساهمة مقفلة')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = BONDS_DARK

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2.paragraph_format.right_to_left = True
    run = subtitle2.add_run('المملكة العربية السعودية')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(14)
    run.font.color.rgb = BONDS_GRAY

    doc.add_paragraph()

    # Contact box on cover - positioned at bottom
    contact_box = doc.add_paragraph()
    contact_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_box.paragraph_format.right_to_left = True
    run = contact_box.add_run(AR_TITLE)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = BONDS_GOLD

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.right_to_left = True
    run = p.add_run(AR_NAME)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(16)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.right_to_left = True
    run = p.add_run(PHONE)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.right_to_left = True
    run = p.add_run(EMAIL)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(12)
    run.font.color.rgb = BONDS_GOLD

    # QR code on cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(qr_path, width=Inches(1.0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.right_to_left = True
    run = p.add_run('www.bonds-global.com')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(10)
    run.font.color.rgb = BONDS_GRAY

    doc.add_page_break()

    # Executive Summary
    add_heading_rtl(doc, 'ملخص تنفيذي', 1)
    add_paragraph_rtl(doc, 'تأسست شركة إحياء الأصول غير المكتملة العقارية كنموذج استثماري متخصص، يهدف إلى إحياء العقارات غير المكتملة في المملكة العربية السعودية من خلال شراكات محاصة مع مالكيها، دون الحاجة إلى شراء الأراضي أو الاعتماد على التمويل الخارجي.')
    add_paragraph_rtl(doc, 'تعمل الشركة بصورة قانونية كنظام مساهمة مقفلة، ممولة بالكامل من المساهمين المؤسسين، وتلتزم بأعلى معايير الكود السعودي للبناء والتشطيب، مع تغطية تأمينية شاملة لجميع المشاريع.')

    add_heading_rtl(doc, 'أبرز المؤشرات المالية', 2)
    rows = [
        ['رأس المال المصرح', '200 مليون ريال'],
        ['إجمالي الاستثمار المستهدف', '226 مليون ريال'],
        ['الاستثمار الفعلي المبدئي', '25 مليون ريال'],
        ['العائد الداخلي المتوقع (IRR)', '19% - 20%'],
        ['مضاعفة رأس المال عند الخروج', '2.5x - 4x'],
        ['فترة استرداد رأس المال', '5.5 سنة'],
    ]
    add_table_rtl(doc, ['المؤشر', 'القيمة'], rows)

    add_heading_rtl(doc, 'الفرصة السوقية', 2)
    add_bullet_rtl(doc, 'السعودية تشهد نمواً عقارياً غير مسبوق بفضل رؤية 2030.')
    add_bullet_rtl(doc, 'آلاف العقارات غير المكتملة متاحة بسبب نقص السيولة أو الخبرة التنفيذية.')
    add_bullet_rtl(doc, 'الطلب على الوحدات الجاهزة يفوق العرض في المدن الرئيسية.')
    add_bullet_rtl(doc, 'عوائد إيجارية مرتفعة في الرياض (8.89%) وجدة (7.89%).')

    add_heading_rtl(doc, 'نموذج العمل', 2)
    add_bullet_rtl(doc, 'الشراكة بالمحاصة مع مالك العقار غير المكتمل.')
    add_bullet_rtl(doc, 'الشركة تتحمل تكاليف التشطيب والتسويق والبيع.')
    add_bullet_rtl(doc, 'المالك يحتفظ بملكية عقاره طوال فترة المشروع.')
    add_bullet_rtl(doc, 'يتم توزيع الأرباح بين الشركة والمالك وفقاً للعقد الموقع.')

    doc.add_page_break()

    # 1. Company Overview
    add_heading_rtl(doc, '1. نظرة عامة على الشركة', 1)
    add_heading_rtl(doc, '1.1 الرؤية', 2)
    add_paragraph_rtl(doc, 'أن نكون الشريك الاستراتيجي المفضل لمالكي العقارات غير المكتملة في المملكة العربية السعودية.')

    add_heading_rtl(doc, '1.2 الرسالة', 2)
    add_paragraph_rtl(doc, 'إحياء العقارات غير المكتملة بمعايير عالية من الجودة والالتزام، وبنموذج مالي مستدام.')

    add_heading_rtl(doc, '1.3 القيم الأساسية', 2)
    add_bullet_rtl(doc, 'الاحترام، الجودة، الشفافية، الاستدامة، الابتكار.')

    add_heading_rtl(doc, '1.4 الشكل القانوني', 2)
    add_paragraph_rtl(doc, 'شركة مساهمة مقفلة مرخصة من وزارة التجارة وهيئة العقار (REGA).')

    add_heading_rtl(doc, '1.5 رأس المال وهيكل الملكية', 2)
    rows = [
        ['رأس المال المصرح', '200,000,000 ريال سعودي'],
        ['عدد الأسهم', '40,000 سهم'],
        ['القيمة الاسمية للسهم', '5,000 ريال'],
        ['القيمة الفعلية للسهم', '5,650 ريال'],
        ['إجمالي الاستثمار المستهدف', '226,000,000 ريال سعودي'],
    ]
    add_table_rtl(doc, ['البند', 'القيمة'], rows)

    doc.add_page_break()

    # 2. Market Analysis
    add_heading_rtl(doc, '2. تحليل السوق', 1)
    add_paragraph_rtl(doc, 'يشهد قطاع الاستثمار والتطوير العقاري في المملكة تحولاً جذرياً مدفوعاً برؤية 2030.')

    add_heading_rtl(doc, '2.1 المؤشرات الرئيسية', 2)
    rows = [
        ['العائد الإيجاري في الرياض', '8.89%'],
        ['العائد الإيجاري في جدة', '7.89%'],
        ['نمو الصفقات العقارية', '+14.2%'],
        ['حجم التمويل العقاري', '730 مليار ريال'],
        ['فجوة السعر/الطلب', '6%'],
    ]
    add_table_rtl(doc, ['المؤشر', 'القيمة'], rows)

    add_heading_rtl(doc, '2.2 المناطق المستهدفة', 2)
    add_paragraph_rtl(doc, 'الرياض: الشمال الفاخر، الشمال الشرقي.')
    add_paragraph_rtl(doc, 'جدة: الساحل الشمالي، الوسط الحيوي.')

    doc.add_page_break()

    # 3. Services
    add_heading_rtl(doc, '3. الخدمات والحلول', 1)
    add_bullet_rtl(doc, 'التقييم الهندسي والقانوني الشامل.')
    add_bullet_rtl(doc, 'إعداد التصاميم المعمارية والإنشائية.')
    add_bullet_rtl(doc, 'إدارة التشطيب والإشراف اليومي.')
    add_bullet_rtl(doc, 'التسويق والبيع الاحترافي.')
    add_bullet_rtl(doc, 'التسليم ونقل الملكية.')

    add_heading_rtl(doc, '3.2 نماذج الشراكة', 2)
    rows = [
        ['النموذج', 'الوصف'],
        ['المحاصة الكاملة', 'تتولى الشركة التشطيب والتسويق مقابل نسبة من الأرباح'],
        ['الشراء الجزئي', 'تشتري الشركة جزءاً من الوحدات بعد التشطيب'],
        ['رسوم الإدارة', 'رسوم ثابتة مقابل التشطيب والتسويق'],
        ['النموذج المختلط', 'مزيج من المحاصة والشراء ورسوم الإدارة'],
    ]
    add_table_rtl(doc, ['النموذج', 'الوصف'], rows)

    doc.add_page_break()

    # 4. Business Model
    add_heading_rtl(doc, '4. النموذج التجاري', 1)
    add_bullet_rtl(doc, 'البحث عن العقارات غير المكتملة.')
    add_bullet_rtl(doc, 'التقييم الهندسي والقانوني والمالي.')
    add_bullet_rtl(doc, 'التعاقد مع المالك.')
    add_bullet_rtl(doc, 'التشطيب وفق الكود السعودي.')
    add_bullet_rtl(doc, 'التسويق والبيع.')
    add_bullet_rtl(doc, 'توزيع الأرباح.')

    add_heading_rtl(doc, '4.2 مصادر الإيرادات', 2)
    rows = [
        ['مصدر الإيراد', 'النسبة'],
        ['المحاصة', '60%'],
        ['إعادة البيع', '25%'],
        ['رسوم الإدارة', '10%'],
        ['خدمات استشارية', '5%'],
    ]
    add_table_rtl(doc, ['المصدر', 'النسبة'], rows)

    doc.add_page_break()

    # 5. Marketing
    add_heading_rtl(doc, '5. استراتيجية التسويق والمبيعات', 1)
    rows = [
        ['القناة', 'الميزانية الشهرية'],
        ['Google Ads', '10,000 ريال'],
        ['Snapchat Ads', '8,000 ريال'],
        ['Instagram / TikTok', '5,000 ريال'],
        ['SEO', '3,000 ريال'],
        ['LinkedIn', '2,000 ريال'],
    ]
    add_table_rtl(doc, ['القناة', 'الميزانية'], rows)

    doc.add_page_break()

    # 6. Operational Plan
    add_heading_rtl(doc, '6. الخطة التشغيلية', 1)
    rows = [
        ['المرحلة', 'المدة'],
        ['التقييم', '4 أسابيع'],
        ['التعاقد', '2-4 أسابيع'],
        ['التصميم', '4-8 أسابيع'],
        ['التشطيب', '16-24 أسبوع'],
        ['التسويق والبيع', '12-24 أسبوع'],
        ['التسليم', '4-8 أسابيع'],
    ]
    add_table_rtl(doc, ['المرحلة', 'المدة'], rows)

    doc.add_page_break()

    # 7. Financial Plan
    add_heading_rtl(doc, '7. الخطة المالية', 1)
    rows = [
        ['المؤشر', 'القيمة'],
        ['IRR', '19% - 20%'],
        ['عائد الخروج', '2.5x - 4x'],
        ['فترة الاسترداد', '5.5 سنة'],
        ['صافي الربح السنة 5', '9.8 مليون ريال'],
    ]
    add_table_rtl(doc, ['المؤشر', 'القيمة'], rows)

    doc.add_page_break()

    # 8. Risk Analysis
    add_heading_rtl(doc, '8. تحليل المخاطر والتحوطات', 1)
    rows = [
        ['المخاطر', 'التحوط'],
        ['عيوب إنشائية مخفية', 'فحص هندسي + تأمين + احتياطي 10%'],
        ['تأخر البيع', 'تسعير تنافسي + عقود مسبقة'],
        ['تأخر المقاول', 'كفالة بنكية + عقوبات'],
        ['نزاعات قانونية', 'عقود واضحة + تحكيم ملزم'],
    ]
    add_table_rtl(doc, ['المخاطر', 'التحوط'], rows)

    doc.add_page_break()

    # 9. Implementation
    add_heading_rtl(doc, '9. خطة التنفيذ والخروج', 1)
    add_bullet_rtl(doc, 'السنة 1-2: إثبات النموذج (3-5 مشاريع).')
    add_bullet_rtl(doc, 'السنة 3-5: النمو والتكرار (6-10 مشاريع سنوياً).')
    add_bullet_rtl(doc, 'السنة 5-7: الخروج الاستراتيجي.')
    add_bullet_rtl(doc, 'خيارات الخروج: بيع حصة مسيطرة، طرح جزئي، إعادة شراء.')

    doc.add_page_break()

    # 10. Conclusion
    add_heading_rtl(doc, '10. الخلاصة والتوصيات', 1)
    add_paragraph_rtl(doc, 'تقدم الشركة فرصة استثمارية فريدة في السوق العقاري السعودي. النموذج يعتمد على الشراكة المحاصة مع مالكي العقارات غير المكتملة، مما يلغي الحاجة إلى شراء الأراضي ويقلل رأس المال المطلوب.')
    add_paragraph_rtl(doc, 'بفضل التمويل الذاتي، والالتزام بالكود السعودي، والتغطية التأمينية الشاملة، تتمتع الشركة بمخاطر محسوبة وعوائد جذابة.')

    doc.save(output_path)
    print(f'Arabic business plan saved: {output_path}')

def create_english_business_plan(logo_path, qr_path, output_path):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_footer_ltr(section)

    # Cover page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(logo_path, width=Inches(2.0))

    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.right_to_left = False
    run = title_p.add_run('Business Plan')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(40)
    run.font.bold = True
    run.font.color.rgb = BONDS_BLUE

    title2_p = doc.add_paragraph()
    title2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title2_p.paragraph_format.right_to_left = False
    run = title2_p.add_run('Real Estate Asset Revitalization Company')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = BONDS_GOLD

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.right_to_left = False
    run = subtitle.add_run('Closed Joint Stock Company')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = BONDS_DARK

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2.paragraph_format.right_to_left = False
    run = subtitle2.add_run('Kingdom of Saudi Arabia')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(14)
    run.font.color.rgb = BONDS_GRAY

    doc.add_paragraph()

    # Contact box on cover
    contact_box = doc.add_paragraph()
    contact_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_box.paragraph_format.right_to_left = False
    run = contact_box.add_run(EN_TITLE)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = BONDS_GOLD

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.right_to_left = False
    run = p.add_run(EN_NAME)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(16)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.right_to_left = False
    run = p.add_run(PHONE)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.right_to_left = False
    run = p.add_run(EMAIL)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(12)
    run.font.color.rgb = BONDS_GOLD

    # QR code
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(qr_path, width=Inches(1.0))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.right_to_left = False
    run = p.add_run('www.bonds-global.com')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(10)
    run.font.color.rgb = BONDS_GRAY

    doc.add_page_break()

    # Executive Summary
    add_heading_ltr(doc, 'Executive Summary', 1)
    add_paragraph_ltr(doc, 'The Real Estate Asset Revitalization Company is a unique investment model aimed at revitalizing incomplete real estate assets in Saudi Arabia through partnership agreements with property owners, without purchasing land or relying on external financing.')
    add_paragraph_ltr(doc, 'The company operates as a closed joint stock company, fully funded by founding shareholders, and adheres to the highest standards of the Saudi Building Code and finishing works, with comprehensive insurance coverage.')

    add_heading_ltr(doc, 'Key Financial Indicators', 2)
    rows = [
        ['Authorized Capital', 'SAR 200 million'],
        ['Total Target Investment', 'SAR 226 million'],
        ['Initial Investment', 'SAR 25 million'],
        ['Expected IRR', '19% - 20%'],
        ['Capital Return at Exit', '2.5x - 4x'],
        ['Payback Period', '5.5 years'],
    ]
    add_table_ltr(doc, ['Indicator', 'Value'], rows)

    add_heading_ltr(doc, 'Market Opportunity', 2)
    add_bullet_ltr(doc, 'Saudi Arabia is witnessing unprecedented real estate growth driven by Vision 2030.')
    add_bullet_ltr(doc, 'Thousands of incomplete assets are available due to lack of liquidity or execution expertise.')
    add_bullet_ltr(doc, 'Demand for ready-built units exceeds supply in major cities.')
    add_bullet_ltr(doc, 'High rental yields in Riyadh (8.89%) and Jeddah (7.89%).')

    add_heading_ltr(doc, 'Business Model', 2)
    add_bullet_ltr(doc, 'Partnership (Muhassah) with asset owners.')
    add_bullet_ltr(doc, 'Company bears finishing, marketing, and selling costs.')
    add_bullet_ltr(doc, 'Owner retains property ownership.')
    add_bullet_ltr(doc, 'Profits shared per signed contract.')

    doc.add_page_break()

    # 1. Company Overview
    add_heading_ltr(doc, '1. Company Overview', 1)
    add_heading_ltr(doc, '1.1 Vision', 2)
    add_paragraph_ltr(doc, 'To be the preferred strategic partner for owners of incomplete real estate assets in Saudi Arabia.')

    add_heading_ltr(doc, '1.2 Mission', 2)
    add_paragraph_ltr(doc, 'Revitalize incomplete real estate assets with high standards of quality and compliance.')

    add_heading_ltr(doc, '1.3 Core Values', 2)
    add_bullet_ltr(doc, 'Respect, Quality, Transparency, Sustainability, Innovation.')

    add_heading_ltr(doc, '1.4 Legal Structure', 2)
    add_paragraph_ltr(doc, 'A closed joint stock company licensed by the Ministry of Commerce and the Real Estate General Authority (REGA).')

    add_heading_ltr(doc, '1.5 Capital Structure', 2)
    rows = [
        ['Authorized Capital', 'SAR 200,000,000'],
        ['Shares', '40,000'],
        ['Nominal Value per Share', 'SAR 5,000'],
        ['Actual Value per Share', 'SAR 5,650'],
        ['Total Target Investment', 'SAR 226,000,000'],
    ]
    add_table_ltr(doc, ['Item', 'Value'], rows)

    doc.add_page_break()

    # 2. Market Analysis
    add_heading_ltr(doc, '2. Market Analysis', 1)
    add_paragraph_ltr(doc, 'The Saudi real estate sector is undergoing radical transformation driven by Vision 2030.')

    add_heading_ltr(doc, '2.1 Key Indicators', 2)
    rows = [
        ['Rental Yield in Riyadh', '8.89%'],
        ['Rental Yield in Jeddah', '7.89%'],
        ['Transaction Growth', '+14.2%'],
        ['Real Estate Financing', 'SAR 730 billion'],
        ['Price/Demand Gap', '6%'],
    ]
    add_table_ltr(doc, ['Indicator', 'Value'], rows)

    add_heading_ltr(doc, '2.2 Target Areas', 2)
    add_paragraph_ltr(doc, 'Riyadh: Premium North, Northeast.')
    add_paragraph_ltr(doc, 'Jeddah: North Coast, Vibrant Center.')

    doc.add_page_break()

    # 3. Services
    add_heading_ltr(doc, '3. Services and Solutions', 1)
    add_bullet_ltr(doc, 'Comprehensive engineering and legal assessment.')
    add_bullet_ltr(doc, 'Architectural and structural design preparation.')
    add_bullet_ltr(doc, 'Finishing management and daily supervision.')
    add_bullet_ltr(doc, 'Professional marketing and sales.')
    add_bullet_ltr(doc, 'Delivery and ownership transfer.')

    add_heading_ltr(doc, '3.2 Partnership Models', 2)
    rows = [
        ['Model', 'Description'],
        ['Full Muhassah', 'Company handles finishing and marketing for profit share'],
        ['Partial Purchase', 'Company buys part of units after finishing'],
        ['Management Fee', 'Fixed fee for finishing and marketing'],
        ['Hybrid Model', 'Mix of Muhassah, purchase, and management fees'],
    ]
    add_table_ltr(doc, ['Model', 'Description'], rows)

    doc.add_page_break()

    # 4. Business Model
    add_heading_ltr(doc, '4. Business Model', 1)
    add_bullet_ltr(doc, 'Identify incomplete real estate assets.')
    add_bullet_ltr(doc, 'Conduct engineering, legal, and financial assessment.')
    add_bullet_ltr(doc, 'Contract with the owner.')
    add_bullet_ltr(doc, 'Finish according to Saudi Building Code.')
    add_bullet_ltr(doc, 'Market and sell.')
    add_bullet_ltr(doc, 'Distribute profits.')

    add_heading_ltr(doc, '4.2 Revenue Sources', 2)
    rows = [
        ['Source', 'Share'],
        ['Muhassah', '60%'],
        ['Resale', '25%'],
        ['Management Fees', '10%'],
        ['Consulting', '5%'],
    ]
    add_table_ltr(doc, ['Source', 'Share'], rows)

    doc.add_page_break()

    # 5. Marketing
    add_heading_ltr(doc, '5. Marketing and Sales Strategy', 1)
    rows = [
        ['Channel', 'Monthly Budget'],
        ['Google Ads', 'SAR 10,000'],
        ['Snapchat Ads', 'SAR 8,000'],
        ['Instagram / TikTok', 'SAR 5,000'],
        ['SEO', 'SAR 3,000'],
        ['LinkedIn', 'SAR 2,000'],
    ]
    add_table_ltr(doc, ['Channel', 'Budget'], rows)

    doc.add_page_break()

    # 6. Operational Plan
    add_heading_ltr(doc, '6. Operational Plan', 1)
    rows = [
        ['Phase', 'Duration'],
        ['Assessment', '4 weeks'],
        ['Contracting', '2-4 weeks'],
        ['Design', '4-8 weeks'],
        ['Finishing', '16-24 weeks'],
        ['Marketing & Sales', '12-24 weeks'],
        ['Delivery', '4-8 weeks'],
    ]
    add_table_ltr(doc, ['Phase', 'Duration'], rows)

    doc.add_page_break()

    # 7. Financial Plan
    add_heading_ltr(doc, '7. Financial Plan', 1)
    rows = [
        ['Indicator', 'Value'],
        ['IRR', '19% - 20%'],
        ['Exit Return', '2.5x - 4x'],
        ['Payback Period', '5.5 years'],
        ['Net Profit Year 5', 'SAR 9.8 million'],
    ]
    add_table_ltr(doc, ['Indicator', 'Value'], rows)

    doc.add_page_break()

    # 8. Risk Analysis
    add_heading_ltr(doc, '8. Risk Analysis and Mitigation', 1)
    rows = [
        ['Risk', 'Mitigation'],
        ['Hidden structural defects', 'Engineering inspection + insurance + 10% reserve'],
        ['Delayed sales', 'Competitive pricing + pre-sale contracts'],
        ['Contractor delay', 'Bank guarantee + penalties'],
        ['Legal disputes', 'Clear contracts + binding arbitration'],
    ]
    add_table_ltr(doc, ['Risk', 'Mitigation'], rows)

    doc.add_page_break()

    # 9. Implementation
    add_heading_ltr(doc, '9. Implementation and Exit', 1)
    add_bullet_ltr(doc, 'Years 1-2: Proof of Concept (3-5 projects).')
    add_bullet_ltr(doc, 'Years 3-5: Growth and Replication (6-10 projects annually).')
    add_bullet_ltr(doc, 'Years 5-7: Strategic Exit.')
    add_bullet_ltr(doc, 'Exit options: Sale of controlling stake, partial IPO, buyback.')

    doc.add_page_break()

    # 10. Conclusion
    add_heading_ltr(doc, '10. Conclusion and Recommendations', 1)
    add_paragraph_ltr(doc, 'The company presents a unique investment opportunity in the Saudi real estate market. The model relies on Muhassah partnerships with owners of incomplete assets, eliminating the need to purchase land and reducing required capital.')
    add_paragraph_ltr(doc, 'With self-funding, compliance with the Saudi Building Code, and comprehensive insurance coverage, the company enjoys calculated risks and attractive returns.')

    doc.save(output_path)
    print(f'English business plan saved: {output_path}')

# Main execution
logo_path = convert_logo()
qr_path = create_qr()

arabic_output = r'C:\Users\vip\bonds-global-web\docs\business-plans\canva-ready-arabic\business-plan-arabic-canva.docx'
english_output = r'C:\Users\vip\bonds-global-web\docs\business-plans\canva-ready-english\business-plan-english-canva.docx'

os.makedirs(os.path.dirname(arabic_output), exist_ok=True)
os.makedirs(os.path.dirname(english_output), exist_ok=True)

create_arabic_business_plan(logo_path, qr_path, arabic_output)
create_english_business_plan(logo_path, qr_path, english_output)

print('All business plans created successfully')
