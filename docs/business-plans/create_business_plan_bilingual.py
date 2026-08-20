from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import os

# Bonds brand colors
BONDS_GOLD = RGBColor(212, 168, 83)
BONDS_DARK = RGBColor(10, 15, 26)
BONDS_GRAY = RGBColor(102, 102, 102)
BONDS_BLUE = RGBColor(0, 51, 102)

def convert_logo():
    logo_webp = r'C:\Users\vip\bonds-global-web\assets\bonds-logo-2026-v2.webp'
    logo_png = r'C:\Users\vip\bonds-global-web\docs\business-plans\bonds-logo.png'
    if not os.path.exists(logo_png):
        img = Image.open(logo_webp)
        img.save(logo_png, 'PNG')
    return logo_png

def set_rtl(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.right_to_left = True

def set_ltr(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.right_to_left = False

def add_heading_rtl(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    set_rtl(heading)
    for run in heading.runs:
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        if level == 1:
            run.font.color.rgb = BONDS_BLUE
            run.font.size = Pt(18)
        elif level == 2:
            run.font.color.rgb = BONDS_GOLD
            run.font.size = Pt(14)
    return heading

def add_heading_ltr(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    set_ltr(heading)
    for run in heading.runs:
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        if level == 1:
            run.font.color.rgb = BONDS_BLUE
            run.font.size = Pt(18)
        elif level == 2:
            run.font.color.rgb = BONDS_GOLD
            run.font.size = Pt(14)
    return heading

def add_paragraph_rtl(doc, text, bold=False, italic=False, size=12, color=RGBColor(0, 0, 0)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    set_rtl(p)
    return p

def add_paragraph_ltr(doc, text, bold=False, italic=False, size=12, color=RGBColor(0, 0, 0)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    set_ltr(p)
    return p

def add_bullet_rtl(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(11)
    set_rtl(p)
    return p

def add_bullet_ltr(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(11)
    set_ltr(p)
    return p

def add_table_rtl(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    table.autofit = True
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            set_rtl(paragraph)
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                run.font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '0A0F1A')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                set_rtl(paragraph)
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    return table

def add_table_ltr(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.autofit = True
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            set_ltr(paragraph)
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                run.font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '0A0F1A')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                set_ltr(paragraph)
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    return table

def add_page_number_rtl(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.right_to_left = True
    run1 = p.add_run('Bonds Financial Management & Consulting  |  ')
    run1.font.name = 'Arial'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run1.font.size = Pt(9)
    run1.font.color.rgb = BONDS_GRAY
    run2 = p.add_run('www.bonds-global.com')
    run2.font.name = 'Arial'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run2.font.size = Pt(9)
    run2.font.color.rgb = BONDS_GOLD

def add_page_number_ltr(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.right_to_left = False
    run1 = p.add_run('Bonds Financial Management & Consulting  |  ')
    run1.font.name = 'Arial'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run1.font.size = Pt(9)
    run1.font.color.rgb = BONDS_GRAY
    run2 = p.add_run('www.bonds-global.com')
    run2.font.name = 'Arial'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run2.font.size = Pt(9)
    run2.font.color.rgb = BONDS_GOLD

def create_arabic_version(logo_path):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_page_number_rtl(section)

    # Cover page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(logo_path, width=Inches(2.5))

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.right_to_left = True
    run = title_p.add_run('خطة عمل')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = BONDS_BLUE

    title2_p = doc.add_paragraph()
    title2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title2_p.paragraph_format.right_to_left = True
    run = title2_p.add_run('شركة إحياء الأصول غير المكتملة العقارية')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = BONDS_GOLD

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.right_to_left = True
    run = subtitle.add_run('شركة مساهمة مقفلة')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = BONDS_DARK

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2.paragraph_format.right_to_left = True
    run = subtitle2.add_run('المملكة العربية السعودية')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(14)
    run.font.color.rgb = BONDS_GRAY

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.right_to_left = True
    run = date_p.add_run('يونيو 2026')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(12)
    run.font.color.rgb = BONDS_GRAY

    conf_p = doc.add_paragraph()
    conf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_p.paragraph_format.right_to_left = True
    run = conf_p.add_run('وثيقة سرية - مخصصة للمساهمين المؤسسين')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = BONDS_GRAY

    doc.add_page_break()

    # Executive Summary
    add_heading_rtl(doc, 'ملخص تنفيذي', 1)
    add_paragraph_rtl(doc, 'تأسست شركة إحياء الأصول غير المكتملة العقارية كنموذج استثماري متخصص، يهدف إلى إحياء العقارات غير المكتملة في المملكة العربية السعودية من خلال شراكات محاصة مع مالكيها، دون الحاجة إلى شراء الأراضي أو الاعتماد على التمويل الخارجي.')
    add_paragraph_rtl(doc, 'تعمل الشركة بصورة قانونية كنظام مساهمة مقفلة، ممولة بالكامل من المساهمين المؤسسين، وتلتزم بأعلى معايير الكود السعودي للبناء والتشطيب، مع تغطية تأمينية شاملة لجميع المشاريع.')

    add_heading_rtl(doc, 'أبرز المؤشرات المالية', 2)
    rows = [
        ['رأس المال المصرح', '200 مليون ريال'],
        ['إجمالي الاستثمار المستهدف', '226 مليون ريال'],
        ['الاستثمار الفعلي المبدئي', '25 مليون ريال'],
        ['العائد الداخلي المتوقع (IRR) على أفق 10 سنوات', '19% - 20%'],
        ['مضاعفة رأس المال عند الخروج', '2.5x - 4x'],
        ['فترة استرداد رأس المال المستدعى', '5.5 سنة'],
    ]
    add_table_rtl(doc, ['المؤشر', 'القيمة'], rows)

    add_heading_rtl(doc, 'الفرصة السوقية', 2)
    add_bullet_rtl(doc, 'السعودية تشهد نمواً عقارياً غير مسبوق بفضل رؤية 2030 ومشاريعها التنموية الضخمة.')
    add_bullet_rtl(doc, 'تتوفر في السوق آلاف العقارات غير المكتملة نتيجة نقص السيولة أو الخبرة التنفيذية.')
    add_bullet_rtl(doc, 'الطلب على الوحدات الجاهزة يفوق العرض في المدن الرئيسية.')
    add_bullet_rtl(doc, 'تسجل الرياض وجدة عوائد إيجارية مرتفعة تصل إلى 8.89% و 7.89% على التوالي.')

    add_heading_rtl(doc, 'نموذج العمل', 2)
    add_bullet_rtl(doc, 'الشراكة بالمحاصة مع مالك العقار غير المكتمل.')
    add_bullet_rtl(doc, 'تتحمل الشركة تكاليف التشطيب والتسويق والبيع.')
    add_bullet_rtl(doc, 'يحتفظ المالك بملكية عقاره طوال فترة المشروع.')
    add_bullet_rtl(doc, 'يتم توزيع الأرباح بين الشركة والمالك وفقاً للعقد الموقع.')

    doc.add_page_break()

    # 1. Company Overview
    add_heading_rtl(doc, '1. نظرة عامة على الشركة', 1)
    add_heading_rtl(doc, '1.1 الرؤية', 2)
    add_paragraph_rtl(doc, 'أن نكون الشريك الاستراتيجي المفضل لمالكي العقارات غير المكتملة في المملكة العربية السعودية، من خلال تقديم حلول احترافية لإحياء أصولهم وتحقيق أقصى قيمة ممكنة لهم وللمساهمين.')

    add_heading_rtl(doc, '1.2 الرسالة', 2)
    add_paragraph_rtl(doc, 'إحياء العقارات غير المكتملة بمعايير عالية من الجودة والالتزام، وبنموذج مالي مستدام يحقق العوائد للشركاء ويحفظ حقوق ملاك الأصول.')

    add_heading_rtl(doc, '1.3 القيم الأساسية', 2)
    add_bullet_rtl(doc, 'الاحترام: احترام مالك الأصل وحقوقه الملكية في كل مرحلة من مراحل المشروع.')
    add_bullet_rtl(doc, 'الجودة: الالتزام بالكود السعودي وأعلى معايير التشطيب والسلامة.')
    add_bullet_rtl(doc, 'الشفافية: وضوح كامل في العقود والتقارير المالية والتشغيلية.')
    add_bullet_rtl(doc, 'الاستدامة: بناء نموذج اقتصادي طويل الأجل يحقق القيمة لجميع الأطراف.')
    add_bullet_rtl(doc, 'الابتكار: استخدام التقنية والحلول الذكية لتحسين الكفاءة وتقليل المخاطر.')

    add_heading_rtl(doc, '1.4 الشكل القانوني', 2)
    add_paragraph_rtl(doc, 'تأسست الشركة بصورة قانونية كنظام مساهمة مقفلة، مرخصة من وزارة التجارة وهيئة العقار (REGA)، وفقاً لنظام الشركات السعودي. يتيح هذا الشكل جمع رأس المال من المساهمين المؤسسين مع الحفاظ على مرونة الإدارة وسرعة اتخاذ القرار.')

    add_heading_rtl(doc, '1.5 رأس المال وهيكل الملكية', 2)
    rows = [
        ['رأس المال المصرح', '200,000,000 ريال سعودي'],
        ['عدد الأسهم', '40,000 سهم'],
        ['القيمة الاسمية للسهم', '5,000 ريال سعودي'],
        ['القيمة الفعلية للسهم', '5,650 ريال سعودي'],
        ['علاوة الإصدار ومصاريف التأسيس', '26,000,000 ريال سعودي'],
        ['إجمالي الاستثمار المستهدف', '226,000,000 ريال سعودي'],
    ]
    add_table_rtl(doc, ['البند', 'القيمة'], rows)

    doc.add_page_break()

    # 2. Market Analysis
    add_heading_rtl(doc, '2. تحليل السوق', 1)
    add_heading_rtl(doc, '2.1 نظرة عامة على السوق العقاري السعودي', 2)
    add_paragraph_rtl(doc, 'يشهد قطاع الاستثمار والتطوير العقاري في المملكة العربية السعودية تحولاً جذرياً، مدفوعاً برؤية 2030 والمشاريع التنموية الضخمة. يتجه القطاع نحو النضج، مع بدء ديناميكيات طلب حقيقية وملموسة في سلوك المستهلكين والمستثمرين.')

    add_heading_rtl(doc, '2.2 المؤشرات الرئيسية', 2)
    rows = [
        ['متوسط العائد الإيجاري في الرياض', '8.89%'],
        ['متوسط العائد الإيجاري في جدة', '7.89%'],
        ['معدل نمو الصفقات العقارية (2025)', '+14.2%'],
        ['حجم التمويل العقاري الإجمالي', '730 مليار ريال'],
        ['متوسط فجوة السعر الفعلي مقارنة بالطلب', '6%'],
        ['علاوة البناء الجديد', '+25%'],
    ]
    add_table_rtl(doc, ['المؤشر', 'القيمة'], rows)

    add_heading_rtl(doc, '2.3 محركات النمو', 2)
    add_bullet_rtl(doc, 'رؤية 2030 والمشاريع العملاقة مثل نيوم، الدرعية، البحر الأحمر، وروشن.')
    add_bullet_rtl(doc, 'برامج التمويل السكني "سكني" الموجهة للوحدات المتوسطة (70 - 120 م²).')
    add_bullet_rtl(doc, 'نمو الرواتب وتوجه الأفراد نحو المساحات الأصغر والأكثر كفاءة.')
    add_bullet_rtl(doc, 'الهجرة الداخلية نحو المدن الكبرى والعاصمة الرياض.')
    add_bullet_rtl(doc, 'انتعاش سوق الإيجارات وارتفاع عوائدها إلى مستويات جاذبة.')

    add_heading_rtl(doc, '2.4 المناطق المستهدفة', 2)
    add_paragraph_rtl(doc, 'مدينة الرياض:', bold=True)
    add_bullet_rtl(doc, 'الشمال الفاخر: العقيق، حطين، الملقا - يجذب التنفيذيين ويوفر عوائد إيجارية مرتفعة.')
    add_bullet_rtl(doc, 'الشمال الشرقي: المونسية، الرمال، الياسمين، النرجس - كثافة سكانية عالية وبنية تحتية حديثة.')

    add_paragraph_rtl(doc, 'مدينة جدة:', bold=True)
    add_bullet_rtl(doc, 'الساحل الشمالي: أبحر الشمالية، الصواري - يمثل مستقبلاً سكنياً وسياحياً واعداً.')
    add_bullet_rtl(doc, 'الوسط الحيوي: الحمراء، النعيم، الزهراء، السلامة - يوفر عوائد إيجارية ثابتة ونشاطاً تجارياً مستمراً.')

    add_heading_rtl(doc, '2.5 تحليل المنافسة', 2)
    add_paragraph_rtl(doc, 'المنافسة الرئيسية ليست من الشركات الكبرى، بل من المطورين التقليديين الذين يتطلب نموذجهم شراء أراضٍ ورأس مال ضخم. تكمن ميزتنا التنافسية في عدم الحاجة إلى شراء الأرض، والتركيز على الأصول غير المكتملة التي لا تجذب المطورين الكبار.')

    doc.add_page_break()

    # 3. Services
    add_heading_rtl(doc, '3. الخدمات والحلول', 1)
    add_heading_rtl(doc, '3.1 الخدمات الأساسية', 2)
    add_bullet_rtl(doc, 'إجراء تقييم هندسي وقانوني شامل للعقار غير المكتمل.')
    add_bullet_rtl(doc, 'إعداد التصاميم المعمارية والإنشائية والميكانيكية المطلوبة.')
    add_bullet_rtl(doc, 'إدارة عمليات التشطيب والإشراف اليومي على التنفيذ.')
    add_bullet_rtl(doc, 'التسويق والبيع الاحترافي للوحدات بأسعار السوق.')
    add_bullet_rtl(doc, 'إدارة عملية التسليم ونقل الملكية وتوزيع الأرباح.')

    add_heading_rtl(doc, '3.2 نماذج الشراكة', 2)
    rows = [
        ['النموذج', 'الوصف', 'الحالة المناسبة'],
        ['المحاصة الكاملة', 'تتولى الشركة التشطيب والتسويق مقابل نسبة متفق عليها من الأرباح', 'المالك يملك سيولة محدودة ويرغب في إحياء العقار'],
        ['الشراء الجزئي', 'تشتري الشركة جزءاً من الوحدات بعد إتمام التشطيب', 'المالك يرغب في تحقيق سيولة فورية'],
        ['رسوم الإدارة', 'تتقاضى الشركة رسوم إدارة ثابتة مقابل التشطيب والتسويق', 'المالك يرغب في الاحتفاظ بمعظم الأرباح'],
        ['النموذج المختلط', 'مزيج متوازن من المحاصة والشراء ورسوم الإدارة', 'معظم الحالات العملية'],
    ]
    add_table_rtl(doc, ['النموذج', 'الوصف', 'الحالة المناسبة'], rows)

    doc.add_page_break()

    # 4. Business Model
    add_heading_rtl(doc, '4. النموذج التجاري', 1)
    add_heading_rtl(doc, '4.1 آلية عمل الشركة', 2)
    add_bullet_rtl(doc, 'البحث عن عقارات غير مكتملة في المواقع المستهدفة وتحديد فرصها.')
    add_bullet_rtl(doc, 'إجراء التقييم الهندسي والقانوني والمالي الشامل.')
    add_bullet_rtl(doc, 'التفاوض مع المالك وإبرام عقد محاصة أو شراكة واضح.')
    add_bullet_rtl(doc, 'إنجاز أعمال التشطيب وفق أعلى المعايير والكود السعودي.')
    add_bullet_rtl(doc, 'تسويق الوحدات وبيعها بأسعار السوق المناسبة.')
    add_bullet_rtl(doc, 'توزيع الأرباح بين الشركة والمالك وفق بنود العقد.')

    add_heading_rtl(doc, '4.2 مصادر الإيرادات', 2)
    rows = [
        ['مصدر الإيراد', 'النسبة المتوقعة'],
        ['إيرادات المحاصة (نصيب الشركة)', '60%'],
        ['أرباح إعادة البيع والشراء', '25%'],
        ['رسوم الإدارة والتسويق', '10%'],
        ['خدمات استشارية وإدارية', '5%'],
    ]
    add_table_rtl(doc, ['مصدر الإيراد', 'النسبة المتوقعة'], rows)

    add_heading_rtl(doc, '4.3 هيكل التكاليف', 2)
    rows = [
        ['البند', 'النسبة من تكلفة المشروع'],
        ['تكاليف التشطيب المباشرة', '55%'],
        ['التراخيص والرسوم البلدية', '4%'],
        ['الاستشارات الهندسية والإشراف', '5%'],
        ['التسويق والمبيعات', '7%'],
        ['احتياطي المخاطر', '10%'],
        ['التكاليف الإدارية التشغيلية', '6%'],
        ['رسوم البيع والنقل', '9%'],
        ['التأمين', '4%'],
    ]
    add_table_rtl(doc, ['البند', 'النسبة'], rows)

    add_heading_rtl(doc, '4.4 نقاط القوة التنافسية', 2)
    add_bullet_rtl(doc, 'لا يتطلب النموذج شراء أراضٍ، مما يقلل رأس المال المطلوب للبدء.')
    add_bullet_rtl(doc, 'سرعة الدخول للسوق مقارنة بالتطوير العقاري التقليدي.')
    add_bullet_rtl(doc, 'شبكة علاقات واسعة مع مالكين، مصفين، محامين، ووسطاء عقاريين.')
    add_bullet_rtl(doc, 'الالتزام بالكود السعودي والتأمين، مما يقلل المخاطر القانونية والمالية.')
    add_bullet_rtl(doc, 'نموذج مالي مرن وممول ذاتياً من المساهمين المؤسسين.')

    doc.add_page_break()

    # 5. Marketing & Sales
    add_heading_rtl(doc, '5. استراتيجية التسويق والمبيعات', 1)
    add_heading_rtl(doc, '5.1 استراتيجية التسويق الرقمي', 2)
    rows = [
        ['القناة', 'الهدف', 'الميزانية الشهرية'],
        ['Google Ads', 'الوصول إلى مالكي العقارات غير المكتملة', '10,000 ريال'],
        ['Snapchat Ads', 'استهداف الفئة العمرية 25 - 45 سنة', '8,000 ريال'],
        ['Instagram / TikTok', 'عرض مشاريع "قبل وبعد"', '5,000 ريال'],
        ['SEO', 'الظهور في نتائج البحث عن إحياء المباني', '3,000 ريال'],
        ['LinkedIn', 'الوصول إلى المستثمرين ومالكي العقارات', '2,000 ريال'],
    ]
    add_table_rtl(doc, ['القناة', 'الهدف', 'الميزانية الشهرية'], rows)

    add_heading_rtl(doc, '5.2 قمع المبيعات', 2)
    add_bullet_rtl(doc, 'الوعي: الوصول إلى 100,000 شخص شهرياً.')
    add_bullet_rtl(doc, 'الاهتمام: تحويل 5,000 شخص إلى مهتمين بالخدمة.')
    add_bullet_rtl(doc, 'التقييم: التواصل مع 500 مالك محتمل شهرياً.')
    add_bullet_rtl(doc, 'النية: عقد 50 اجتماعاً مبدئياً شهرياً.')
    add_bullet_rtl(doc, 'الشراء: إبرام 2 - 4 عقود محاصة شهرياً.')

    add_heading_rtl(doc, '5.3 استراتيجية الوصول إلى المالكين', 2)
    add_bullet_rtl(doc, 'بناء قاعدة بيانات شاملة للعقارات غير المكتملة في المناطق المستهدفة.')
    add_bullet_rtl(doc, 'إبرام شراكات استراتيجية مع المصفين والوسطاء العقاريين.')
    add_bullet_rtl(doc, 'المشاركة في المعارض العقارية السنوية داخل المملكة.')
    add_bullet_rtl(doc, 'إطلاق حملات توعوية حول فوائد إحياء الأصول غير المكتملة.')

    doc.add_page_break()

    # 6. Operational Plan
    add_heading_rtl(doc, '6. الخطة التشغيلية', 1)
    add_heading_rtl(doc, '6.1 مراحل المشروع الواحد', 2)
    rows = [
        ['المرحلة', 'المدة', 'الأنشطة الرئيسية'],
        ['التقييم', '4 أسابيع', 'فحص هيكلي، مراجعة قانونية، تقدير تكلفة مبدئي'],
        ['التعاقد', '2 - 4 أسابيع', 'صياغة عقد المحاصة، استصدار الموافقات اللازمة'],
        ['التصميم', '4 - 8 أسابيع', 'إعداد التصاميم واستصدار التراخيص البلدية'],
        ['التشطيب', '16 - 24 أسبوع', 'اختيار مقاول، إشراف يومي، تفتيش دوري'],
        ['التسويق والبيع', '12 - 24 أسبوع', 'عرض الوحدات، استقبال العملاء، عقود البيع'],
        ['التسليم والتسوية', '4 - 8 أسابيع', 'تسليم الوحدات، نقل الملكية، توزيع الأرباح'],
    ]
    add_table_rtl(doc, ['المرحلة', 'المدة', 'الأنشطة الرئيسية'], rows)

    add_heading_rtl(doc, '6.2 مؤشرات الأداء الرئيسية (KPIs)', 2)
    rows = [
        ['المؤشر', 'الهدف', 'التحذير', 'الخطر'],
        ['نسبة إنجاز التشطيب', 'أكثر من 80% في الشهر الرابع', 'أقل من 60%', 'تأخر المقاول'],
        ['معدل بيع الوحدات', 'أكثر من 30% في الشهر الثاني', 'أقل من 10%', 'ضعف الطلب'],
        ['تكلفة التشطيب مقارنة بالمتوقع', 'أقل من 105%', 'أكثر من 115%', 'عيوب مخفية'],
        ['رضا العملاء (NPS)', 'أكثر من 50', 'أقل من 30', 'مشاكل في الجودة'],
        ['مدة البيع', 'أقل من 6 أشهر', 'أكثر من 9 أشهر', 'ركود في السوق'],
    ]
    add_table_rtl(doc, ['المؤشر', 'الهدف', 'التحذير', 'الخطر'], rows)

    add_heading_rtl(doc, '6.3 نظام الجودة', 2)
    add_bullet_rtl(doc, 'الالتزام الكامل بكود البناء السعودي في جميع مراحل التنفيذ.')
    add_bullet_rtl(doc, 'إجراء فحوصات جودة دورية في كل مرحلة من مراحل التشطيب.')
    add_bullet_rtl(doc, 'توثيق جميع المواد المستخدمة وشهادات المطابقة الخاصة بها.')
    add_bullet_rtl(doc, 'تسليم المشروع بملف فني وقانوني كامل للمالك والمشتري.')

    doc.add_page_break()

    # 7. Management Team
    add_heading_rtl(doc, '7. الفريق الإداري والتنظيمي', 1)
    add_heading_rtl(doc, '7.1 الهيكل التنظيمي', 2)
    add_paragraph_rtl(doc, 'تم تصميم الهيكل الإداري ليكون مرناً وفعالاً، مع بدء تشغيلي يعتمد على 15 موظفاً في السنوات الأولى، مع إمكانية التوسع تبعاً لنمو المحفظة.')

    add_heading_rtl(doc, '7.2 الأدوار الإدارية الرئيسية', 2)
    rows = [
        ['الدور', 'المسؤوليات الرئيسية'],
        ['الرئيس التنفيذي', 'وضع الاستراتيجية العامة، إدارة علاقات المستثمرين، اتخاذ القرارات الكبرى'],
        ['المدير التشغيلي', 'إدارة المشاريع، التنسيق مع المقاولين، ضمان جودة التنفيذ'],
        ['المدير المالي', 'إدارة التدفقات النقدية، إعداد التقارير المالية، التخطيط المالي'],
        ['المدير التجاري', 'إدارة التسويق والمبيعات، بناء علاقات المالكين، تطوير الأعمال'],
        ['المدير القانوني', 'صياغة العقود، ضمان الامتثال، حل النزاعات'],
        ['مدير الموارد البشرية والإدارة', 'إدارة الموظفين والإجراءات الإدارية'],
    ]
    add_table_rtl(doc, ['الدور', 'المسؤوليات الرئيسية'], rows)

    add_heading_rtl(doc, '7.3 خطة التوظيف', 2)
    rows = [
        ['السنة', 'عدد الموظفين', 'الملاحظات'],
        ['السنة الأولى', '10 - 15', 'فريق أساسي لتشغيل المرحلة الأولى'],
        ['السنة الثانية', '15', 'تثبيت الهيكل التشغيلي'],
        ['السنة الثالثة', '18 - 20', 'التوسع في عدد المشاريع'],
        ['السنة الخامسة', '25 - 30', 'دعم محفظة 10 مشاريع سنوياً'],
    ]
    add_table_rtl(doc, ['السنة', 'عدد الموظفين', 'الملاحظات'], rows)

    doc.add_page_break()

    # 8. Financial Plan
    add_heading_rtl(doc, '8. الخطة المالية', 1)
    add_heading_rtl(doc, '8.1 افتراضات النموذج المالي', 2)
    rows = [
        ['الافتراض', 'القيمة'],
        ['متوسط مساحة المبنى', '1,500 متر مربع'],
        ['عدد الوحدات في المبنى', '6 - 8 وحدات'],
        ['تكلفة التشطيب', '1,200 - 1,800 ريال لكل متر مربع'],
        ['سعر البيع', '2,500 - 3,500 ريال لكل متر مربع'],
        ['نسبة المحاصة للشركة', '50%'],
    ]
    add_table_rtl(doc, ['الافتراض', 'القيمة'], rows)

    add_heading_rtl(doc, '8.2 استخدامات رأس المال', 2)
    rows = [
        ['الاستخدام', 'المبلغ', 'النسبة'],
        ['استثمارات في مشاريع (شراء + تشطيب)', '150 مليون ريال', '66%'],
        ['رأس المال العامل التشغيلي', '40 مليون ريال', '18%'],
        ['تكنولوجيا وأنظمة', '10 مليون ريال', '4%'],
        ['تأسيس وتسويق وعلاقات', '15 مليون ريال', '7%'],
        ['احتياطي سيولة طوارئ', '11 مليون ريال', '5%'],
    ]
    add_table_rtl(doc, ['الاستخدام', 'المبلغ', 'النسبة'], rows)

    add_heading_rtl(doc, '8.3 ملخص النتائج المالية المتوقعة', 2)
    rows = [
        ['المؤشر', 'القيمة'],
        ['العائد الداخلي المتوقع (IRR) على أفق 10 سنوات', '19% - 20%'],
        ['مضاعفة رأس المال عند الخروج', '2.5x - 4x'],
        ['فترة استرداد رأس المال المستدعى', '5.5 سنة'],
        ['صافي الربح المتوقع في السنة الخامسة', '9.8 مليون ريال'],
        ['القيمة السوقية المتوقعة بعد 5 سنوات', '120 - 155 مليون ريال'],
    ]
    add_table_rtl(doc, ['المؤشر', 'القيمة'], rows)

    add_heading_rtl(doc, '8.4 آلية سحب رأس المال', 2)
    add_paragraph_rtl(doc, 'لا يُضخ رأس المال دفعة واحدة، بل يتم سحبه على أربع دفعات مرتبطة بإنجاز مراحل محددة:')
    add_bullet_rtl(doc, 'الدفعة الأولى (30%): التأسيس وشراء 3 - 5 مبانٍ أولى.')
    add_bullet_rtl(doc, 'الدفعة الثانية (30%): إنجاز أول مشروعين وتحقيق أولى الإيرادات.')
    add_bullet_rtl(doc, 'الدفعة الثالثة (25%): الوصول إلى 6 مشاريع نشطة وتحقيق التعادل التشغيلي.')
    add_bullet_rtl(doc, 'الدفعة الرابعة (15%): احتياطي مرن يُطلق حسب الحاجة أو يُعاد للمساهمين.')

    doc.add_page_break()

    # 9. Risk Analysis
    add_heading_rtl(doc, '9. تحليل المخاطر والتحوطات', 1)
    add_heading_rtl(doc, '9.1 المخاطر والتحوطات المقترحة', 2)
    rows = [
        ['المخاطر', 'الاحتمالية', 'التأثير', 'التحوط'],
        ['صعوبة الحصول على أصول جيدة', 'عالية', 'توقف النمو', 'بناء شبكة علاقات وقاعدة بيانات وشراكات مع مصفين'],
        ['عيوب إنشائية مخفية', 'عالية', 'خسارة مالية كبيرة', 'فحص هندسي شامل + تأمين + احتياطي 10%'],
        ['تأخر بيع الوحدات', 'متوسطة', 'حمل تكاليف إضافي', 'تسعير تنافسي + عقود مسبقة البيع'],
        ['تأخر المقاول في التنفيذ', 'متوسطة', 'تكاليف إضافية', 'كفالة بنكية + عقوبات تأخير في العقد'],
        ['تغير أسعار مواد البناء', 'متوسطة', 'ارتفاع التكاليف', 'عقود ثابتة السعر مع مقاولين موثوقين'],
        ['نزاعات قانونية مع المالك', 'متوسطة', 'توقف المشروع', 'عقود واضحة + إدراج شرط تحكيم ملزم'],
        ['تراجع أسعار العقار', 'منخفضة إلى متوسطة', 'انخفاض الهوامش', 'تنويع جغرافي واستهداف فئات متعددة'],
    ]
    add_table_rtl(doc, ['المخاطر', 'الاحتمالية', 'التأثير', 'التحوط'], rows)

    add_heading_rtl(doc, '9.2 سياسة التأمين', 2)
    add_bullet_rtl(doc, 'تأمين المقاولين الشامل (CAR) لكل مشروع على حدة.')
    add_bullet_rtl(doc, 'تأمين المسؤولية المدنية العامة للشركة.')
    add_bullet_rtl(doc, 'تأمين العيوب الخفية للمباني غير المكتملة.')
    add_bullet_rtl(doc, 'تأمين المسؤولية المهنية للمهندسين والاستشاريين.')
    add_bullet_rtl(doc, 'تأمين مقر الشركة والأصول الثابتة.')

    add_heading_rtl(doc, '9.3 الامتثال للكود السعودي', 2)
    add_bullet_rtl(doc, 'كود البناء السعودي (SBC) - معايير الإنشاء والسلامة.')
    add_bullet_rtl(doc, 'اشتراطات البلدية - التراخيص والتفتيش الدوري.')
    add_bullet_rtl(doc, 'اشتراطات الدفاع المدني - السلامة من الحريق.')
    add_bullet_rtl(doc, 'اشتراطات الهيئة العامة للعقار (REGA).')

    doc.add_page_break()

    # 10. Implementation Milestones
    add_heading_rtl(doc, '10. خطة التنفيذ والمراحل الزمنية', 1)
    add_heading_rtl(doc, '10.1 المرحلة الأولى: إثبات النموذج (السنة 1 - 2)', 2)
    add_bullet_rtl(doc, 'إتمام التأسيس القانوني واستخراج التراخيص اللازمة.')
    add_bullet_rtl(doc, 'ضخ 25 مليون ريال كاستثمار أولي للمرحلة التجريبية.')
    add_bullet_rtl(doc, 'تنفيذ 3 - 5 مشاريع قياسية لإثبات جدوى النموذج.')
    add_bullet_rtl(doc, 'بناء فريق تشغيلي أساسي يتراوح بين 10 و 15 موظفاً.')
    add_bullet_rtl(doc, 'إنشاء قاعدة بيانات شاملة للأصول غير المكتملة.')

    add_heading_rtl(doc, '10.2 المرحلة الثانية: النمو والتكرار (السنة 3 - 5)', 2)
    add_bullet_rtl(doc, 'زيادة رأس المال المستدعى تدريجياً حسب نمو المحفظة.')
    add_bullet_rtl(doc, 'الوصول إلى 6 - 10 مشاريع سنوياً.')
    add_bullet_rtl(doc, 'تحقيق الربحية التشغيلية المستدامة.')
    add_bullet_rtl(doc, 'تطوير نظام ERP وأتمتة العمليات التشغيلية.')
    add_bullet_rtl(doc, 'التركيز على مدينتي الرياض وجدة كأسواق رئيسية.')

    add_heading_rtl(doc, '10.3 المرحلة الثالثة: الخروج الاستراتيجي (السنة 5 - 7)', 2)
    add_bullet_rtl(doc, 'بيع حصة مسيطرة لصندوق استثماري عقاري أو مطور كبير.')
    add_bullet_rtl(doc, 'أو إجراء طرح جزئي في السوق المالية السعودية (Tadawul).')
    add_bullet_rtl(doc, 'أو إعادة شراء الأسهم من المساهمين تدريجياً من الأرباح.')

    add_heading_rtl(doc, '10.4 المخطط الزمني الرئيسي', 2)
    rows = [
        ['المرحلة', 'المدة', 'المعالم الرئيسية'],
        ['التأسيس', '0 - 6 أشهر', 'استخراج التراخيص، تجهيز المقر، تجميع الفريق الأساسي'],
        ['الإطلاق', '6 - 12 شهراً', 'تنفيذ أول 3 مشاريع، تحقيق أولى الإيرادات'],
        ['النمو', 'السنة 2 - 3', '6 مشاريع سنوياً، الوصول للتعادل التشغيلي'],
        ['التوسع', 'السنة 4 - 5', '10 مشاريع سنوياً، أرباح مستدامة'],
        ['الخروج', 'السنة 5 - 7', 'بيع استراتيجي أو طرح جزئي'],
    ]
    add_table_rtl(doc, ['المرحلة', 'المدة', 'المعالم الرئيسية'], rows)

    doc.add_page_break()

    # 11. Conclusion
    add_heading_rtl(doc, '11. الخلاصة والتوصيات', 1)
    add_paragraph_rtl(doc, 'تقدم شركة إحياء الأصول غير المكتملة العقارية فرصة استثمارية فريدة في السوق العقاري السعودي. يعتمد النموذج التجاري على الشراكة المحاصة مع مالكي العقارات غير المكتملة، مما يلغي الحاجة إلى شراء الأراضي ويقلل رأس المال المطلوب للبدء.')
    add_paragraph_rtl(doc, 'بفضل التمويل الذاتي من المساهمين المؤسسين، والالتزام بالكود السعودي، والتغطية التأمينية الشاملة، تتمتع الشركة بمخاطر محسوبة وعوائد جذابة تصل إلى 19 - 20% كعائد داخلي متوقع على أفق 10 سنوات.')

    add_heading_rtl(doc, 'التوصيات الرئيسية', 2)
    add_bullet_rtl(doc, 'البدء بمرحلة تجريبية بـ 25 مليون ريال وتنفيذ 3 - 5 مشاريع قياسية.')
    add_bullet_rtl(doc, 'التركيز على مدينة واحدة (الرياض أو جدة) في السنوات الأولى لتجنب تشتيت الجهد.')
    add_bullet_rtl(doc, 'بناء قاعدة بيانات للأصول غير المكتملة قبل الإطلاق الرسمي.')
    add_bullet_rtl(doc, 'إبرام شراكات مع مقاولين معتمدين يقدمون كفالات بنكية على أعمالهم.')
    add_bullet_rtl(doc, 'إدراج شرط تحكيم ملزم في جميع عقود المحاصة.')
    add_bullet_rtl(doc, 'تطوير نظام ERP لإدارة المحفظة والتدفقات النقدية.')
    add_bullet_rtl(doc, 'إعداد عقد محاصة موحد يحفظ حقوق الشركة والمالك على حد سواء.')

    doc.save('C:/Users/vip/bonds-global-web/docs/business-plans/business-plan-arabic-final.docx')
    print('Arabic business plan with logo created successfully')

def create_english_version(logo_path):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_page_number_ltr(section)

    # Cover page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(logo_path, width=Inches(2.5))

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.right_to_left = False
    run = title_p.add_run('Business Plan')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = BONDS_BLUE

    title2_p = doc.add_paragraph()
    title2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title2_p.paragraph_format.right_to_left = False
    run = title2_p.add_run('Real Estate Asset Revitalization Company')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = BONDS_GOLD

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.right_to_left = False
    run = subtitle.add_run('Closed Joint Stock Company')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = BONDS_DARK

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2.paragraph_format.right_to_left = False
    run = subtitle2.add_run('Kingdom of Saudi Arabia')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(14)
    run.font.color.rgb = BONDS_GRAY

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.right_to_left = False
    run = date_p.add_run('June 2026')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(12)
    run.font.color.rgb = BONDS_GRAY

    conf_p = doc.add_paragraph()
    conf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_p.paragraph_format.right_to_left = False
    run = conf_p.add_run('Confidential - For Founding Shareholders Only')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = BONDS_GRAY

    doc.add_paragraph()
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.right_to_left = False
    run = contact_p.add_run('Financial Advisor')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = BONDS_GOLD

    contact_p2 = doc.add_paragraph()
    contact_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p2.paragraph_format.right_to_left = False
    run = contact_p2.add_run('Dr. Talal bin Hassan Al-Zahrani')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(16)
    run.font.bold = True

    contact_p3 = doc.add_paragraph()
    contact_p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p3.paragraph_format.right_to_left = False
    run = contact_p3.add_run('+966 56 756 6616')
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(12)
    run.font.color.rgb = BONDS_GRAY

    doc.add_page_break()

    # Executive Summary
    add_heading_ltr(doc, 'Executive Summary', 1)
    add_paragraph_ltr(doc, 'The Real Estate Asset Revitalization Company is established as a unique investment model aimed at revitalizing incomplete real estate assets in the Kingdom of Saudi Arabia through partnership agreements with property owners, without the need to purchase land or rely on external financing.')
    add_paragraph_ltr(doc, 'The company operates legally as a closed joint stock company, fully funded by its founding shareholders, and adheres to the highest standards of the Saudi Building Code and finishing works, with comprehensive insurance coverage for all projects.')

    add_heading_ltr(doc, 'Key Financial Indicators', 2)
    rows = [
        ['Authorized Capital', 'SAR 200 million'],
        ['Total Target Investment', 'SAR 226 million'],
        ['Initial Actual Investment', 'SAR 25 million'],
        ['Expected IRR over 10 Years', '19% - 20%'],
        ['Capital Return Multiple at Exit', '2.5x - 4x'],
        ['Payback Period', '5.5 years'],
    ]
    add_table_ltr(doc, ['Indicator', 'Value'], rows)

    add_heading_ltr(doc, 'Market Opportunity', 2)
    add_bullet_ltr(doc, 'Saudi Arabia is witnessing unprecedented real estate growth driven by Vision 2030 and its massive development projects.')
    add_bullet_ltr(doc, 'Thousands of incomplete real estate assets are available due to lack of liquidity or execution expertise.')
    add_bullet_ltr(doc, 'Demand for ready-built units exceeds supply in major cities.')
    add_bullet_ltr(doc, 'Riyadh and Jeddah record high rental yields reaching 8.89% and 7.89% respectively.')

    add_heading_ltr(doc, 'Business Model', 2)
    add_bullet_ltr(doc, 'Partnership (Muhassah) with the owner of the incomplete asset.')
    add_bullet_ltr(doc, 'The company bears finishing, marketing, and selling costs.')
    add_bullet_ltr(doc, 'The owner retains property ownership throughout the project.')
    add_bullet_ltr(doc, 'Profits are distributed between the company and the owner according to the signed contract.')

    doc.add_page_break()

    # 1. Company Overview
    add_heading_ltr(doc, '1. Company Overview', 1)
    add_heading_ltr(doc, '1.1 Vision', 2)
    add_paragraph_ltr(doc, 'To be the preferred strategic partner for owners of incomplete real estate assets in the Kingdom of Saudi Arabia, by providing professional solutions to revitalize their assets and maximize value for them and the shareholders.')

    add_heading_ltr(doc, '1.2 Mission', 2)
    add_paragraph_ltr(doc, 'Revitalize incomplete real estate assets with high standards of quality and compliance, through a sustainable financial model that generates returns for partners while preserving the rights of asset owners.')

    add_heading_ltr(doc, '1.3 Core Values', 2)
    add_bullet_ltr(doc, 'Respect: Respecting the asset owner and his property rights at every stage of the project.')
    add_bullet_ltr(doc, 'Quality: Compliance with the Saudi Building Code and highest finishing and safety standards.')
    add_bullet_ltr(doc, 'Transparency: Full clarity in contracts, financial reports, and operational reporting.')
    add_bullet_ltr(doc, 'Sustainability: Building a long-term economic model that creates value for all parties.')
    add_bullet_ltr(doc, 'Innovation: Using technology and smart solutions to improve efficiency and reduce risks.')

    add_heading_ltr(doc, '1.4 Legal Structure', 2)
    add_paragraph_ltr(doc, 'The company is established as a closed joint stock company, licensed by the Ministry of Commerce and the Real Estate General Authority (REGA), in accordance with the Saudi Companies Law. This structure allows raising capital from founding shareholders while maintaining management flexibility and fast decision-making.')

    add_heading_ltr(doc, '1.5 Capital and Ownership Structure', 2)
    rows = [
        ['Authorized Capital', 'SAR 200,000,000'],
        ['Number of Shares', '40,000 shares'],
        ['Nominal Value per Share', 'SAR 5,000'],
        ['Actual Value per Share', 'SAR 5,650'],
        ['Issue Premium and Establishment Costs', 'SAR 26,000,000'],
        ['Total Target Investment', 'SAR 226,000,000'],
    ]
    add_table_ltr(doc, ['Item', 'Value'], rows)

    doc.add_page_break()

    # 2. Market Analysis
    add_heading_ltr(doc, '2. Market Analysis', 1)
    add_heading_ltr(doc, '2.1 Overview of the Saudi Real Estate Market', 2)
    add_paragraph_ltr(doc, 'The real estate investment and development sector in the Kingdom of Saudi Arabia is undergoing a radical transformation, driven by Vision 2030 and massive development projects. The sector is maturing, with real demand dynamics becoming tangible in consumer and investor behavior.')

    add_heading_ltr(doc, '2.2 Key Indicators', 2)
    rows = [
        ['Average Rental Yield in Riyadh', '8.89%'],
        ['Average Rental Yield in Jeddah', '7.89%'],
        ['Real Estate Transaction Growth (2025)', '+14.2%'],
        ['Total Real Estate Financing Volume', 'SAR 730 billion'],
        ['Average Price-to-Demand Gap', '6%'],
        ['New Construction Premium', '+25%'],
    ]
    add_table_ltr(doc, ['Indicator', 'Value'], rows)

    add_heading_ltr(doc, '2.3 Growth Drivers', 2)
    add_bullet_ltr(doc, 'Vision 2030 and mega projects such as NEOM, Diriyah, Red Sea Project, and Roshn.')
    add_bullet_ltr(doc, '"Sakani" housing finance programs targeting medium-sized units (70 - 120 m²).')
    add_bullet_ltr(doc, 'Salary growth and trend toward smaller, more efficient spaces.')
    add_bullet_ltr(doc, 'Internal migration toward major cities and the capital Riyadh.')
    add_bullet_ltr(doc, 'Recovery of the rental market and attractive rental yields.')

    add_heading_ltr(doc, '2.4 Target Areas', 2)
    add_paragraph_ltr(doc, 'Riyadh:', bold=True)
    add_bullet_ltr(doc, 'Premium North: Al-Aqeeq, Hittin, Al-Malqa - attracts executives and offers high rental yields.')
    add_bullet_ltr(doc, 'Northeast: Al-Munsiyah, Al-Rimal, Al-Yasmin, Al-Narjis - high population density and modern infrastructure.')

    add_paragraph_ltr(doc, 'Jeddah:', bold=True)
    add_bullet_ltr(doc, 'North Coast: Abhur Al-Shamaliyah, Al-Sawari - promising residential and tourism future.')
    add_bullet_ltr(doc, 'Vibrant Central: Al-Hamra, Al-Naeem, Al-Zahra, Al-Salama - stable rental yields and continuous commercial activity.')

    add_heading_ltr(doc, '2.5 Competitive Analysis', 2)
    add_paragraph_ltr(doc, 'The main competition is not from large corporations, but from traditional developers whose model requires land purchase and large capital. Our competitive advantage lies in not needing to buy land, and focusing on incomplete assets that do not attract large developers.')

    doc.add_page_break()

    # 3. Services
    add_heading_ltr(doc, '3. Services and Solutions', 1)
    add_heading_ltr(doc, '3.1 Core Services', 2)
    add_bullet_ltr(doc, 'Comprehensive engineering and legal assessment of incomplete assets.')
    add_bullet_ltr(doc, 'Preparation of architectural, structural, and mechanical designs.')
    add_bullet_ltr(doc, 'Management of finishing operations and daily supervision.')
    add_bullet_ltr(doc, 'Professional marketing and sales of units at market prices.')
    add_bullet_ltr(doc, 'Management of delivery, ownership transfer, and profit distribution.')

    add_heading_ltr(doc, '3.2 Partnership Models', 2)
    rows = [
        ['Model', 'Description', 'Suitable For'],
        ['Full Muhassah', 'Company handles finishing and marketing for an agreed profit share', 'Owner with limited liquidity seeking asset revitalization'],
        ['Partial Purchase', 'Company buys part of units after finishing', 'Owner seeking immediate liquidity'],
        ['Management Fee', 'Company charges fixed management fee for finishing and marketing', 'Owner wishing to retain most profits'],
        ['Hybrid Model', 'Balanced mix of Muhassah, purchase, and management fees', 'Most practical cases'],
    ]
    add_table_ltr(doc, ['Model', 'Description', 'Suitable For'], rows)

    doc.add_page_break()

    # 4. Business Model
    add_heading_ltr(doc, '4. Business Model', 1)
    add_heading_ltr(doc, '4.1 How the Company Works', 2)
    add_bullet_ltr(doc, 'Identify incomplete real estate assets in target locations.')
    add_bullet_ltr(doc, 'Conduct comprehensive engineering, legal, and financial assessment.')
    add_bullet_ltr(doc, 'Negotiate with the owner and sign a clear Muhassah or partnership contract.')
    add_bullet_ltr(doc, 'Execute finishing works according to highest standards and Saudi Building Code.')
    add_bullet_ltr(doc, 'Market and sell units at appropriate market prices.')
    add_bullet_ltr(doc, 'Distribute profits between the company and owner per contract terms.')

    add_heading_ltr(doc, '4.2 Revenue Sources', 2)
    rows = [
        ['Revenue Source', 'Expected Share'],
        ['Muhassah Revenues (Company Share)', '60%'],
        ['Resale and Purchase Profits', '25%'],
        ['Management and Marketing Fees', '10%'],
        ['Consulting and Administrative Services', '5%'],
    ]
    add_table_ltr(doc, ['Revenue Source', 'Expected Share'], rows)

    add_heading_ltr(doc, '4.3 Cost Structure', 2)
    rows = [
        ['Item', 'Share of Project Cost'],
        ['Direct Finishing Costs', '55%'],
        ['Municipal Licenses and Fees', '4%'],
        ['Engineering Consultation and Supervision', '5%'],
        ['Marketing and Sales', '7%'],
        ['Risk Reserve', '10%'],
        ['Operational Administrative Costs', '6%'],
        ['Sales and Transfer Fees', '9%'],
        ['Insurance', '4%'],
    ]
    add_table_ltr(doc, ['Item', 'Share'], rows)

    add_heading_ltr(doc, '4.4 Competitive Strengths', 2)
    add_bullet_ltr(doc, 'No land purchase required, reducing initial capital needs.')
    add_bullet_ltr(doc, 'Faster market entry compared to traditional real estate development.')
    add_bullet_ltr(doc, 'Extensive network with owners, appraisers, lawyers, and real estate brokers.')
    add_bullet_ltr(doc, 'Compliance with Saudi Building Code and insurance, reducing legal and financial risks.')
    add_bullet_ltr(doc, 'Flexible financial model self-funded by founding shareholders.')

    doc.add_page_break()

    # 5. Marketing & Sales
    add_heading_ltr(doc, '5. Marketing and Sales Strategy', 1)
    add_heading_ltr(doc, '5.1 Digital Marketing Strategy', 2)
    rows = [
        ['Channel', 'Objective', 'Monthly Budget'],
        ['Google Ads', 'Reach owners of incomplete real estate assets', 'SAR 10,000'],
        ['Snapchat Ads', 'Target age group 25 - 45', 'SAR 8,000'],
        ['Instagram / TikTok', 'Showcase "before and after" projects', 'SAR 5,000'],
        ['SEO', 'Appear in search results for building revitalization', 'SAR 3,000'],
        ['LinkedIn', 'Reach investors and property owners', 'SAR 2,000'],
    ]
    add_table_ltr(doc, ['Channel', 'Objective', 'Monthly Budget'], rows)

    add_heading_ltr(doc, '5.2 Sales Funnel', 2)
    add_bullet_ltr(doc, 'Awareness: Reach 100,000 people monthly.')
    add_bullet_ltr(doc, 'Interest: Convert 5,000 people into interested prospects.')
    add_bullet_ltr(doc, 'Evaluation: Contact 500 potential owners monthly.')
    add_bullet_ltr(doc, 'Intent: Hold 50 initial meetings monthly.')
    add_bullet_ltr(doc, 'Purchase: Sign 2 - 4 Muhassah contracts monthly.')

    add_heading_ltr(doc, '5.3 Owner Outreach Strategy', 2)
    add_bullet_ltr(doc, 'Build a comprehensive database of incomplete real estate assets in target areas.')
    add_bullet_ltr(doc, 'Establish strategic partnerships with appraisers and real estate brokers.')
    add_bullet_ltr(doc, 'Participate in annual real estate exhibitions within the Kingdom.')
    add_bullet_ltr(doc, 'Launch awareness campaigns about the benefits of revitalizing incomplete assets.')

    doc.add_page_break()

    # 6. Operational Plan
    add_heading_ltr(doc, '6. Operational Plan', 1)
    add_heading_ltr(doc, '6.1 Project Phases', 2)
    rows = [
        ['Phase', 'Duration', 'Key Activities'],
        ['Assessment', '4 weeks', 'Structural inspection, legal review, initial cost estimate'],
        ['Contracting', '2 - 4 weeks', 'Draft Muhassah contract, obtain necessary approvals'],
        ['Design', '4 - 8 weeks', 'Prepare designs and obtain municipal licenses'],
        ['Finishing', '16 - 24 weeks', 'Select contractor, daily supervision, periodic inspection'],
        ['Marketing and Sales', '12 - 24 weeks', 'Display units, receive clients, sales contracts'],
        ['Delivery and Settlement', '4 - 8 weeks', 'Deliver units, transfer ownership, distribute profits'],
    ]
    add_table_ltr(doc, ['Phase', 'Duration', 'Key Activities'], rows)

    add_heading_ltr(doc, '6.2 Key Performance Indicators (KPIs)', 2)
    rows = [
        ['Indicator', 'Target', 'Warning', 'Risk'],
        ['Finishing Completion Rate', '>80% by month 4', '<60%', 'Contractor delay'],
        ['Unit Sales Rate', '>30% by month 2', '<10%', 'Weak demand'],
        ['Finishing Cost vs. Budget', '<105%', '>115%', 'Hidden defects'],
        ['Customer Satisfaction (NPS)', '>50', '<30', 'Quality issues'],
        ['Sales Duration', '<6 months', '>9 months', 'Market slowdown'],
    ]
    add_table_ltr(doc, ['Indicator', 'Target', 'Warning', 'Risk'], rows)

    add_heading_ltr(doc, '6.3 Quality System', 2)
    add_bullet_ltr(doc, 'Full compliance with the Saudi Building Code at all execution stages.')
    add_bullet_ltr(doc, 'Periodic quality inspections at every phase of finishing.')
    add_bullet_ltr(doc, 'Documentation of all materials used and their compliance certificates.')
    add_bullet_ltr(doc, 'Project handover with complete technical and legal file for owner and buyer.')

    doc.add_page_break()

    # 7. Management Team
    add_heading_ltr(doc, '7. Management Team and Organizational Structure', 1)
    add_heading_ltr(doc, '7.1 Organizational Structure', 2)
    add_paragraph_ltr(doc, 'The organizational structure is designed to be flexible and efficient, with an initial operational team of 15 employees in the early years, with the ability to expand according to portfolio growth.')

    add_heading_ltr(doc, '7.2 Key Management Roles', 2)
    rows = [
        ['Role', 'Key Responsibilities'],
        ['Chief Executive Officer', 'Overall strategy, investor relations, major decision-making'],
        ['Chief Operating Officer', 'Project management, contractor coordination, quality assurance'],
        ['Chief Financial Officer', 'Cash flow management, financial reporting, financial planning'],
        ['Commercial Director', 'Marketing and sales, owner relationships, business development'],
        ['Legal Director', 'Contract drafting, compliance, dispute resolution'],
        ['HR and Administration Manager', 'Employee management and administrative procedures'],
    ]
    add_table_ltr(doc, ['Role', 'Key Responsibilities'], rows)

    add_heading_ltr(doc, '7.3 Hiring Plan', 2)
    rows = [
        ['Year', 'Number of Employees', 'Notes'],
        ['Year 1', '10 - 15', 'Core team for initial phase operation'],
        ['Year 2', '15', 'Stabilize operational structure'],
        ['Year 3', '18 - 20', 'Expand number of projects'],
        ['Year 5', '25 - 30', 'Support portfolio of 10 projects annually'],
    ]
    add_table_ltr(doc, ['Year', 'Number of Employees', 'Notes'], rows)

    doc.add_page_break()

    # 8. Financial Plan
    add_heading_ltr(doc, '8. Financial Plan', 1)
    add_heading_ltr(doc, '8.1 Financial Model Assumptions', 2)
    rows = [
        ['Assumption', 'Value'],
        ['Average Building Area', '1,500 m²'],
        ['Number of Units per Building', '6 - 8 units'],
        ['Finishing Cost', 'SAR 1,200 - 1,800 per m²'],
        ['Selling Price', 'SAR 2,500 - 3,500 per m²'],
        ['Company Muhassah Share', '50%'],
    ]
    add_table_ltr(doc, ['Assumption', 'Value'], rows)

    add_heading_ltr(doc, '8.2 Capital Utilization', 2)
    rows = [
        ['Use of Funds', 'Amount', 'Share'],
        ['Project Investments (Purchase + Finishing)', 'SAR 150 million', '66%'],
        ['Working Capital', 'SAR 40 million', '18%'],
        ['Technology and Systems', 'SAR 10 million', '4%'],
        ['Establishment, Marketing, and Relations', 'SAR 15 million', '7%'],
        ['Emergency Liquidity Reserve', 'SAR 11 million', '5%'],
    ]
    add_table_ltr(doc, ['Use of Funds', 'Amount', 'Share'], rows)

    add_heading_ltr(doc, '8.3 Summary of Expected Financial Results', 2)
    rows = [
        ['Indicator', 'Value'],
        ['Expected IRR over 10 Years', '19% - 20%'],
        ['Capital Return Multiple at Exit', '2.5x - 4x'],
        ['Payback Period', '5.5 years'],
        ['Expected Net Profit in Year 5', 'SAR 9.8 million'],
        ['Expected Market Value after 5 Years', 'SAR 120 - 155 million'],
    ]
    add_table_ltr(doc, ['Indicator', 'Value'], rows)

    add_heading_ltr(doc, '8.4 Capital Drawdown Mechanism', 2)
    add_paragraph_ltr(doc, 'Capital is not injected in a single payment, but drawn in four phases linked to achieving specific milestones:')
    add_bullet_ltr(doc, 'First Tranche (30%): Establishment and purchase of 3 - 5 initial buildings.')
    add_bullet_ltr(doc, 'Second Tranche (30%): Completion of first two projects and realization of initial revenues.')
    add_bullet_ltr(doc, 'Third Tranche (25%): Reaching 6 active projects and achieving operational break-even.')
    add_bullet_ltr(doc, 'Fourth Tranche (15%): Flexible reserve to be released as needed or returned to shareholders.')

    doc.add_page_break()

    # 9. Risk Analysis
    add_heading_ltr(doc, '9. Risk Analysis and Mitigation', 1)
    add_heading_ltr(doc, '9.1 Risks and Mitigation Measures', 2)
    rows = [
        ['Risk', 'Probability', 'Impact', 'Mitigation'],
        ['Difficulty acquiring quality assets', 'High', 'Growth stagnation', 'Build network, database, and partnerships with appraisers'],
        ['Hidden structural defects', 'High', 'Major financial loss', 'Thorough engineering inspection + insurance + 10% reserve'],
        ['Delayed unit sales', 'Medium', 'Carrying cost burden', 'Competitive pricing + pre-sale contracts'],
        ['Contractor delay', 'Medium', 'Additional costs', 'Bank guarantee + contractual delay penalties'],
        ['Construction material price changes', 'Medium', 'Cost increases', 'Fixed-price contracts with reliable contractors'],
        ['Legal disputes with owners', 'Medium', 'Project stoppage', 'Clear contracts + binding arbitration clause'],
        ['Real estate price decline', 'Low to Medium', 'Margin compression', 'Geographic diversification and multiple target segments'],
    ]
    add_table_ltr(doc, ['Risk', 'Probability', 'Impact', 'Mitigation'], rows)

    add_heading_ltr(doc, '9.2 Insurance Policy', 2)
    add_bullet_ltr(doc, 'Contractors All Risk (CAR) insurance for each project.')
    add_bullet_ltr(doc, 'General public liability insurance for the company.')
    add_bullet_ltr(doc, 'Hidden defects insurance for incomplete buildings.')
    add_bullet_ltr(doc, 'Professional liability insurance for engineers and consultants.')
    add_bullet_ltr(doc, 'Insurance for company headquarters and fixed assets.')

    add_heading_ltr(doc, '9.3 Compliance with Saudi Codes', 2)
    add_bullet_ltr(doc, 'Saudi Building Code (SBC) - construction and safety standards.')
    add_bullet_ltr(doc, 'Municipal requirements - licensing and periodic inspection.')
    add_bullet_ltr(doc, 'Civil Defense requirements - fire safety.')
    add_bullet_ltr(doc, 'Real Estate General Authority (REGA) requirements.')

    doc.add_page_break()

    # 10. Implementation Milestones
    add_heading_ltr(doc, '10. Implementation Plan and Timeline', 1)
    add_heading_ltr(doc, '10.1 Phase 1: Proof of Concept (Year 1 - 2)', 2)
    add_bullet_ltr(doc, 'Complete legal establishment and obtain necessary licenses.')
    add_bullet_ltr(doc, 'Inject SAR 25 million as initial investment for the pilot phase.')
    add_bullet_ltr(doc, 'Execute 3 - 5 benchmark projects to prove model viability.')
    add_bullet_ltr(doc, 'Build a core operational team of 10 - 15 employees.')
    add_bullet_ltr(doc, 'Establish a comprehensive database of incomplete assets.')

    add_heading_ltr(doc, '10.2 Phase 2: Growth and Replication (Year 3 - 5)', 2)
    add_bullet_ltr(doc, 'Gradually increase called-up capital according to portfolio growth.')
    add_bullet_ltr(doc, 'Reach 6 - 10 projects annually.')
    add_bullet_ltr(doc, 'Achieve sustainable operational profitability.')
    add_bullet_ltr(doc, 'Develop ERP system and automate operational processes.')
    add_bullet_ltr(doc, 'Focus on Riyadh and Jeddah as primary markets.')

    add_heading_ltr(doc, '10.3 Phase 3: Strategic Exit (Year 5 - 7)', 2)
    add_bullet_ltr(doc, 'Sell controlling stake to a real estate investment fund or major developer.')
    add_bullet_ltr(doc, 'Or conduct partial IPO on the Saudi Stock Exchange (Tadawul).')
    add_bullet_ltr(doc, 'Or gradually buy back shares from shareholders using profits.')

    add_heading_ltr(doc, '10.4 Main Timeline', 2)
    rows = [
        ['Phase', 'Duration', 'Key Milestones'],
        ['Establishment', '0 - 6 months', 'Obtain licenses, set up office, assemble core team'],
        ['Launch', '6 - 12 months', 'Execute first 3 projects, achieve initial revenues'],
        ['Growth', 'Year 2 - 3', '6 projects annually, reach operational break-even'],
        ['Expansion', 'Year 4 - 5', '10 projects annually, sustainable profits'],
        ['Exit', 'Year 5 - 7', 'Strategic sale or partial IPO'],
    ]
    add_table_ltr(doc, ['Phase', 'Duration', 'Key Milestones'], rows)

    doc.add_page_break()

    # 11. Conclusion
    add_heading_ltr(doc, '11. Conclusion and Recommendations', 1)
    add_paragraph_ltr(doc, 'The Real Estate Asset Revitalization Company offers a unique investment opportunity in the Saudi real estate market. The business model relies on Muhassah partnerships with owners of incomplete assets, eliminating the need to purchase land and reducing the initial capital required.')
    add_paragraph_ltr(doc, 'Thanks to self-funding by founding shareholders, compliance with the Saudi Building Code, and comprehensive insurance coverage, the company enjoys calculated risks and attractive returns reaching 19 - 20% expected IRR over a 10-year horizon.')

    add_heading_ltr(doc, 'Key Recommendations', 2)
    add_bullet_ltr(doc, 'Start with a pilot phase of SAR 25 million and execute 3 - 5 benchmark projects.')
    add_bullet_ltr(doc, 'Focus on one city (Riyadh or Jeddah) in the early years to avoid spreading efforts.')
    add_bullet_ltr(doc, 'Build a database of incomplete assets before official launch.')
    add_bullet_ltr(doc, 'Partner with certified contractors who provide bank guarantees for their work.')
    add_bullet_ltr(doc, 'Include binding arbitration clauses in all Muhassah contracts.')
    add_bullet_ltr(doc, 'Develop an ERP system to manage the portfolio and cash flows.')
    add_bullet_ltr(doc, 'Prepare a standardized Muhassah contract that protects both the company and the owner.')

    doc.save('C:/Users/vip/bonds-global-web/docs/business-plans/business-plan-english.docx')
    print('English business plan with logo created successfully')

# Main execution
logo_path = convert_logo()
create_arabic_version(logo_path)
create_english_version(logo_path)
print('All business plans created successfully')
