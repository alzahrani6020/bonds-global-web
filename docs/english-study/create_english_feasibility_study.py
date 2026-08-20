from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_ltr(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.right_to_left = False

def add_heading_ltr(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    set_ltr(heading)
    for run in heading.runs:
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        if level == 1:
            run.font.color.rgb = RGBColor(0, 51, 102)
        elif level == 2:
            run.font.color.rgb = RGBColor(212, 168, 83)
    return heading

def add_paragraph_ltr(doc, text, bold=False, italic=False, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    set_ltr(p)
    return p

def add_bullet_ltr(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(11)
    set_ltr(p)
    return p

def add_table_ltr(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            set_ltr(paragraph)
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                run.font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = __import__('docx.oxml', fromlist=['OxmlElement']).OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '0A0F1A')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                set_ltr(paragraph)
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    return table

# Create document
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

# Cover Page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Feasibility Study')
run.font.size = Pt(44)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Real Estate Asset Revitalization Company')
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(212, 168, 83)
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Closed Joint Stock Company')
run.font.size = Pt(18)
run.font.bold = True
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Kingdom of Saudi Arabia')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(102, 102, 102)
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Prepared by: Bonds Financial Management & Consulting')
run.font.size = Pt(12)
run.font.italic = True
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('June 2026')
run.font.size = Pt(12)
run.font.italic = True
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Financial Advisor')
run.font.size = Pt(14)
run.font.bold = True
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Dr. Talal bin Hassan Al-Zahrani')
run.font.size = Pt(16)
run.font.bold = True
run.font.color.rgb = RGBColor(212, 168, 83)
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('+966 56 756 6616')
run.font.size = Pt(14)
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

doc.add_page_break()

# Table of Contents
add_heading_ltr(doc, 'Table of Contents', 1)
items = [
    'Introduction: Investment Climate and Purchasing Power',
    '1. Executive Summary',
    '2. Project Overview',
    '3. Market Analysis',
    '4. Commercial Model',
    '5. Integrated Business Model',
    '6. Organizational Structure and Team',
    '7. Operational Plan',
    '8. Financial Tables',
    '9. Detailed Financial Tables',
    '10. Break-even and Sensitivity Analysis',
    '11. Risk Analysis and Mitigation',
    '12. Three Scenarios',
    '13. Marketing and Sales Plan',
    '14. Operations and Quality Plan',
    '15. Conclusion and Recommendations',
    '16. Market Enrichments',
    '17. References and Sources',
    'Appendix A: Independent Strategic Analysis',
    'Appendix B: Investor Evaluation',
]
for item in items:
    add_bullet_ltr(doc, item)

doc.add_page_break()

# Introduction
add_heading_ltr(doc, 'Introduction: Investment Climate and Purchasing Power', 1)
add_paragraph_ltr(doc, 'Under the umbrella of Vision 2030, the real estate investment and development sector in the Kingdom of Saudi Arabia is witnessing a radical transformation. The sector is becoming one of the most vital and mature sectors in the region, driven by real demand dynamics rather than speculative randomness.')
add_paragraph_ltr(doc, 'This study was prepared with extreme care based on official data from the Saudi Central Bank (SAMA), the Real Estate General Authority (REGA), the Ministry of Justice, and analyses from Knight Frank and KAPSARC. The goal is to present a realistic market vision that enables investors and building owners to make informed decisions.')

add_heading_ltr(doc, 'Key Market Indicators', 2)
rows = [
    ['Rental Yield in Riyadh', '8.89%'],
    ['Rental Yield in Jeddah', '7.89%'],
    ['Real Estate Transaction Growth (2025)', '14.2%+'],
    ['Real Estate Financing Volume', 'SAR 730 billion'],
    ['Capital Return at Exit', '2.5x - 4x'],
]
add_table_ltr(doc, ['Indicator', 'Value'], rows)

add_paragraph_ltr(doc, 'While rental yield represents annual operating income, the primary return for investors comes from capital appreciation upon strategic sale after revitalizing the assets. Over a 10-year horizon, the expected IRR is approximately 18.5%, with a capital return multiple of 2.5x - 4x at exit in years 5-7.')

add_heading_ltr(doc, 'Growth Drivers', 2)
add_bullet_ltr(doc, 'Mega projects such as NEOM, Red Sea Project, Diriyah, and Roshn.')
add_bullet_ltr(doc, 'Sakani financing programs for medium-sized units (70-120 m²).')
add_bullet_ltr(doc, 'Salary growth and migration toward major cities.')
add_bullet_ltr(doc, 'Stable and attractive rental yields.')

doc.add_page_break()

# 1. Executive Summary
add_heading_ltr(doc, '1. Executive Summary', 1)
add_heading_ltr(doc, '1.1 Core Concept', 2)
add_paragraph_ltr(doc, 'The Real Estate Asset Revitalization Company was founded on a vision to revitalize incomplete buildings without purchasing land. The model operates through partnership (Muhassah) contracts with building owners, where the company covers finishing, marketing, and selling costs, then shares profits with the original owner.')

add_heading_ltr(doc, 'Key Strengths', 2)
add_bullet_ltr(doc, 'No land purchase required - reducing market entry risks and capital needs.')
add_bullet_ltr(doc, 'The Saudi market contains thousands of incomplete buildings.')
add_bullet_ltr(doc, 'Healthy profit margins of 8-12% compared to traditional real estate development.')
add_bullet_ltr(doc, 'Year 1 is foundational; real revenues begin from Year 2.')
add_bullet_ltr(doc, 'Initial required capital is only SAR 25 million, with a flexible ceiling of SAR 226 million.')
add_bullet_ltr(doc, 'True value lies in accumulated asset portfolio, brand, and recurring revenue.')
add_bullet_ltr(doc, 'Primary return comes from capital appreciation at strategic exit.')

add_heading_ltr(doc, 'Financial Highlights', 2)
rows = [
    ['Authorized Capital', 'SAR 200,000,000'],
    ['Shares', '40,000'],
    ['Nominal Value per Share', 'SAR 5,000'],
    ['Actual Value per Share', 'SAR 5,650'],
    ['Total Target Investment', 'SAR 226,000,000'],
    ['Expected IRR (10 years)', '18.5% - 20%'],
    ['Capital Return at Exit', '2.5x - 4x'],
    ['Payback Period', '6 years'],
]
add_table_ltr(doc, ['Indicator', 'Value'], rows)

doc.add_page_break()

# 2. Project Overview
add_heading_ltr(doc, '2. Project Overview', 1)
add_heading_ltr(doc, '2.1 Legal Structure', 2)
add_paragraph_ltr(doc, 'A closed joint stock company licensed by the Ministry of Commerce and the Real Estate General Authority (REGA), in accordance with the Saudi Companies Law. This structure enables capital raising from founding shareholders while maintaining management flexibility and decision-making speed.')

add_heading_ltr(doc, '2.2 Capital Structure', 2)
rows = [
    ['Item', 'Amount (SAR)'],
    ['Authorized Capital', '200,000,000'],
    ['Number of Shares', '40,000'],
    ['Nominal Value per Share', '5,000'],
    ['Actual Value per Share', '5,650'],
    ['Issue Premium and Establishment Costs', '26,000,000'],
    ['Total Target Investment', '226,000,000'],
]
add_table_ltr(doc, ['Item', 'Amount (SAR)'], rows)

add_heading_ltr(doc, '2.3 Capital Utilization', 2)
rows = [
    ['Use of Funds', 'Amount (SAR million)', 'Percentage'],
    ['Project Investments (Purchase + Finishing)', '150', '66%'],
    ['Working Capital', '40', '18%'],
    ['Technology and Systems', '10', '4%'],
    ['Establishment, Marketing, and Relations', '15', '7%'],
    ['Emergency Liquidity Reserve', '11', '5%'],
]
add_table_ltr(doc, ['Use of Funds', 'Amount', 'Share'], rows)

add_heading_ltr(doc, '2.4 Phased Capital Drawdown', 2)
add_bullet_ltr(doc, 'First Tranche (30%): Establishment and purchase of first 3-5 buildings.')
add_bullet_ltr(doc, 'Second Tranche (30%): Completion of first projects and initial revenue.')
add_bullet_ltr(doc, 'Third Tranche (25%): Reaching 6 active projects and operational break-even.')
add_bullet_ltr(doc, 'Fourth Tranche (15%): Flexible reserve for future needs or shareholder distribution.')

doc.add_page_break()

# 3. Market Analysis
add_heading_ltr(doc, '3. Market Analysis', 1)
add_paragraph_ltr(doc, 'The Saudi real estate market is experiencing structural transformation. Major cities like Riyadh and Jeddah offer exceptional rental yields compared to global capitals. The demand for ready-built units exceeds supply, creating a significant opportunity for asset revitalization.')

add_heading_ltr(doc, '3.1 Target Areas', 2)
add_paragraph_ltr(doc, 'Riyadh:', bold=True)
add_bullet_ltr(doc, 'Premium North: Al-Aqeeq, Hittin, Al-Malqa - high rental yields and executive demand.')
add_bullet_ltr(doc, 'Northeast: Al-Munsiyah, Al-Rimal, Al-Yasmin, Al-Narjis - high density and modern infrastructure.')

add_paragraph_ltr(doc, 'Jeddah:', bold=True)
add_bullet_ltr(doc, 'North Coast: Abhur Al-Shamaliyah, Al-Sawari - long-term residential and tourism potential.')
add_bullet_ltr(doc, 'Vibrant Center: Al-Hamra, Al-Naeem, Al-Zahra, Al-Salama - stable rental yields.')

add_heading_ltr(doc, '3.2 Competitive Advantage', 2)
add_paragraph_ltr(doc, 'Traditional developers require land purchase and large capital. Our model focuses on incomplete assets through partnership, allowing faster execution, lower risk, and access to opportunities that large developers overlook.')

doc.add_page_break()

# 4. Commercial Model
add_heading_ltr(doc, '4. Commercial Model', 1)
add_paragraph_ltr(doc, 'The company operates through a hybrid partnership model that balances risk and return:')

add_heading_ltr(doc, 'Partnership Models', 2)
rows = [
    ['Model', 'Description'],
    ['Full Muhassah', 'Company handles finishing and marketing for agreed profit share'],
    ['Partial Purchase', 'Company buys part of units after finishing'],
    ['Management Fee', 'Fixed fee for finishing and marketing services'],
    ['Hybrid Model', 'Combination of Muhassah, purchase, and management fees'],
]
add_table_ltr(doc, ['Model', 'Description'], rows)

add_heading_ltr(doc, 'Revenue Sources', 2)
rows = [
    ['Source', 'Expected Share'],
    ['Muhassah Revenues', '60%'],
    ['Resale Profits', '25%'],
    ['Management Fees', '10%'],
    ['Consulting Services', '5%'],
]
add_table_ltr(doc, ['Source', 'Share'], rows)

doc.add_page_break()

# 5. Integrated Business Model
add_heading_ltr(doc, '5. Integrated Business Model', 1)
add_paragraph_ltr(doc, 'The integrated business model combines multiple revenue streams and cost efficiencies to create a sustainable and scalable operation. By leveraging technology, professional networks, and standardized processes, the company can replicate success across multiple projects.')

add_bullet_ltr(doc, 'Centralized project assessment and legal review.')
add_bullet_ltr(doc, 'Standardized finishing specifications and contractor management.')
add_bullet_ltr(doc, 'Unified marketing and sales platform.')
add_bullet_ltr(doc, 'Data-driven asset selection and pricing.')

doc.add_page_break()

# 6. Organizational Structure
add_heading_ltr(doc, '6. Organizational Structure and Team', 1)
add_paragraph_ltr(doc, 'The organizational structure is designed to be lean and efficient, starting with 15 employees in the early years.')

add_heading_ltr(doc, 'Key Roles', 2)
rows = [
    ['Role', 'Responsibilities'],
    ['CEO', 'Strategy, investor relations, major decisions'],
    ['COO', 'Project management, contractor coordination, quality assurance'],
    ['CFO', 'Cash flow management, financial reporting, planning'],
    ['Commercial Director', 'Marketing, sales, owner relationships'],
    ['Legal Director', 'Contracts, compliance, dispute resolution'],
    ['HR & Admin Manager', 'Employees and administrative procedures'],
]
add_table_ltr(doc, ['Role', 'Responsibilities'], rows)

add_heading_ltr(doc, 'Hiring Plan', 2)
rows = [
    ['Year', 'Employees', 'Notes'],
    ['1', '10 - 15', 'Core team for initial phase'],
    ['2', '15', 'Stabilize operations'],
    ['3', '18 - 20', 'Expand project volume'],
    ['5', '25 - 30', 'Support 10 projects annually'],
]
add_table_ltr(doc, ['Year', 'Employees', 'Notes'], rows)

doc.add_page_break()

# 7. Operational Plan
add_heading_ltr(doc, '7. Operational Plan', 1)
add_heading_ltr(doc, '7.1 Project Phases', 2)
rows = [
    ['Phase', 'Duration', 'Activities'],
    ['Assessment', '4 weeks', 'Structural inspection, legal review, cost estimate'],
    ['Contracting', '2 - 4 weeks', 'Muhassah contract, approvals'],
    ['Design', '4 - 8 weeks', 'Architectural design, municipal licenses'],
    ['Finishing', '16 - 24 weeks', 'Contractor selection, daily supervision'],
    ['Marketing & Sales', '12 - 24 weeks', 'Unit display, client reception, sales contracts'],
    ['Delivery & Settlement', '4 - 8 weeks', 'Unit delivery, ownership transfer, profit sharing'],
]
add_table_ltr(doc, ['Phase', 'Duration', 'Activities'], rows)

add_heading_ltr(doc, '7.2 Project KPIs', 2)
rows = [
    ['KPI', 'Target', 'Warning', 'Risk'],
    ['Finishing Completion', '>80% by month 4', '<60%', 'Contractor delay'],
    ['Unit Sales Rate', '>30% by month 2', '<10%', 'Weak demand'],
    ['Finishing Cost vs Budget', '<105%', '>115%', 'Hidden defects'],
    ['Customer NPS', '>50', '<30', 'Quality issues'],
    ['Sales Duration', '<6 months', '>9 months', 'Market slowdown'],
]
add_table_ltr(doc, ['KPI', 'Target', 'Warning', 'Risk'], rows)

doc.add_page_break()

# 8. Financial Tables
add_heading_ltr(doc, '8. Financial Tables', 1)
add_heading_ltr(doc, '8.1 Financial Model Assumptions', 2)
rows = [
    ['Assumption', 'Value'],
    ['Average Building Area', '1,500 m²'],
    ['Units per Building', '6 - 8 units'],
    ['Finishing Cost', 'SAR 1,200 - 1,800 per m²'],
    ['Selling Price', 'SAR 2,500 - 3,500 per m²'],
    ['Company Muhassah Share', '50%'],
]
add_table_ltr(doc, ['Assumption', 'Value'], rows)

add_heading_ltr(doc, '8.2 Estimated Cost per Project', 2)
rows = [
    ['Item', 'Cost (SAR)', 'Percentage'],
    ['Finishing Costs', '2,250,000', '55%'],
    ['Licenses and Fees', '150,000', '4%'],
    ['Engineering Consultation', '200,000', '5%'],
    ['Marketing and Sales', '300,000', '7%'],
    ['Risk Reserve (15%)', '450,000', '11%'],
    ['Operational Costs', '250,000', '6%'],
    ['Financial Costs', '200,000', '5%'],
    ['Sales and Transfer Fees', '375,000', '9%'],
    ['Total', '4,175,000', '100%'],
]
add_table_ltr(doc, ['Item', 'Cost', 'Share'], rows)

add_heading_ltr(doc, '8.3 Five-Year Income Statement (Hybrid Model)', 2)
rows = [
    ['Item', 'Year 1', 'Year 2', 'Year 3', 'Year 4', 'Year 5'],
    ['Total Revenue (SAR million)', '3.2', '6.4', '9.6', '12.8', '16.0'],
    ['Direct Costs', '-0.4', '-0.8', '-1.2', '-1.6', '-2.0'],
    ['G&A Costs', '-3.2', '-3.2', '-3.5', '-3.8', '-4.0'],
    ['Net Profit before Zakat', '-0.5', '2.4', '5.0', '7.7', '10.4'],
]
add_table_ltr(doc, ['Item', 'Y1', 'Y2', 'Y3', 'Y4', 'Y5'], rows)

doc.add_page_break()

# 9. Detailed Financial Tables
add_heading_ltr(doc, '9. Detailed Financial Tables', 1)
add_paragraph_ltr(doc, 'The detailed financial model includes monthly cash flow for the first year, project cost breakdowns, and return metrics. The model assumes a gradual ramp-up with 2-3 projects in Year 1, reaching 10 projects by Year 5.')

add_heading_ltr(doc, '9.1 Return Metrics', 2)
rows = [
    ['Metric', 'Value'],
    ['Capital Return at Exit (Years 5-7)', '2.5x - 4x'],
    ['Expected IRR (10 years)', '18.5%+'],
    ['Expected Market Value (Year 5)', 'SAR 75 - 94 million'],
    ['Payback Period', '6 years'],
]
add_table_ltr(doc, ['Metric', 'Value'], rows)

add_heading_ltr(doc, '9.2 Long-Term Projection (Years 6-10)', 2)
rows = [
    ['Year', 'Projects', 'Revenue (SAR million)', 'Net Profit (SAR million)'],
    ['6', '12', '19.2', '6.0+'],
    ['7', '14', '22.4', '8.0+'],
    ['8', '16', '25.6', '10.0+'],
    ['9', '18', '28.8', '12.0+'],
    ['10', '20', '32.0', '14.0+'],
]
add_table_ltr(doc, ['Year', 'Projects', 'Revenue', 'Net Profit'], rows)

doc.add_page_break()

# 10. Break-even and Sensitivity
add_heading_ltr(doc, '10. Break-even and Sensitivity Analysis', 1)
add_heading_ltr(doc, '10.1 Break-even Point', 2)
add_paragraph_ltr(doc, 'Under the hybrid model, the company reaches operational break-even with approximately 5-6 projects annually. This assumes a balanced mix of Muhassah, partial purchase, and management fee revenues.')

add_heading_ltr(doc, '10.2 Sensitivity Analysis', 2)
rows = [
    ['Price Change', 'Year 3 Net Profit', 'Year 5 Net Profit'],
    ['-20%', '1.5M', '5.5M'],
    ['-10%', '3.2M', '7.8M'],
    ['Base (0%)', '5.0M', '10.4M'],
    ['+10%', '6.8M', '13.0M'],
    ['+20%', '8.5M', '15.5M'],
]
add_table_ltr(doc, ['Price Change', 'Y3 Profit', 'Y5 Profit'], rows)

add_paragraph_ltr(doc, 'The model shows resilience even with 20% price decline or 30% cost increase, remaining profitable from Year 3 onward.')

doc.add_page_break()

# 11. Risk Analysis
add_heading_ltr(doc, '11. Risk Analysis and Mitigation', 1)
rows = [
    ['Risk', 'Probability', 'Impact', 'Mitigation'],
    ['Difficulty acquiring quality assets', 'High', 'Growth stagnation', 'Network, database, appraisers'],
    ['Hidden structural defects', 'High', 'Major financial loss', 'Engineering inspection + insurance + 15% reserve'],
    ['Delayed unit sales', 'Medium', 'Carrying costs', 'Competitive pricing + pre-sale contracts'],
    ['Contractor delay', 'Medium', 'Additional costs', 'Bank guarantee + penalties'],
    ['Material price changes', 'Medium', 'Cost increases', 'Fixed-price contracts'],
    ['Legal disputes with owners', 'Medium', 'Project stoppage', 'Clear contracts + binding arbitration'],
    ['Real estate price decline', 'Low-Medium', 'Margin compression', 'Geographic diversification'],
]
add_table_ltr(doc, ['Risk', 'Probability', 'Impact', 'Mitigation'], rows)

add_heading_ltr(doc, 'Recommended Risk Reserve', 2)
rows = [
    ['Reserve Type', 'Amount'],
    ['Structural Defects Reserve', 'SAR 225,000 per project'],
    ['Sales Delay Reserve', 'SAR 150,000 per project'],
    ['General Reserve', 'SAR 2,500,000'],
    ['Total Recommended Reserve', 'SAR 3,000,000'],
]
add_table_ltr(doc, ['Reserve', 'Amount'], rows)

doc.add_page_break()

# 12. Three Scenarios
add_heading_ltr(doc, '12. Three Scenarios', 1)
rows = [
    ['Scenario', 'Year 1 Projects', 'Finishing Cost', 'Sales Duration', 'Year 5 Net Profit'],
    ['Pessimistic', '1', '+20% above expected', '+12 months', 'SAR 3.0M'],
    ['Base Case', '2 - 3', 'As expected', '3 - 6 months', 'SAR 10.4M'],
    ['Optimistic', '4', '-10% below expected', '3 months', 'SAR 14.0M'],
]
add_table_ltr(doc, ['Scenario', 'Y1 Projects', 'Cost', 'Sales', 'Y5 Profit'], rows)

add_heading_ltr(doc, 'Break-even by Scenario', 2)
rows = [
    ['Scenario', 'Annual Break-even', 'Cumulative Break-even'],
    ['Pessimistic', '18 projects', '54 projects'],
    ['Base Case', '6 projects', '20 projects'],
    ['Optimistic', '10 projects', '20 projects'],
]
add_table_ltr(doc, ['Scenario', 'Annual', 'Cumulative'], rows)

doc.add_page_break()

# 13. Marketing and Sales Plan
add_heading_ltr(doc, '13. Marketing and Sales Plan', 1)
add_heading_ltr(doc, '13.1 Digital Marketing Strategy', 2)
rows = [
    ['Channel', 'Objective', 'Monthly Budget (SAR)'],
    ['Google Ads', 'Reach owners of incomplete buildings', '10,000'],
    ['Snapchat Ads', 'Target age group 25-45', '8,000'],
    ['Instagram / TikTok', 'Showcase before/after projects', '5,000'],
    ['SEO', 'Rank for "building revitalization"', '3,000'],
    ['LinkedIn', 'Reach investors and owners', '2,000'],
]
add_table_ltr(doc, ['Channel', 'Objective', 'Budget'], rows)

add_heading_ltr(doc, '13.2 Sales Funnel', 2)
add_bullet_ltr(doc, 'Awareness: 100,000 people/month')
add_bullet_ltr(doc, 'Interest: 5,000 people/month')
add_bullet_ltr(doc, 'Evaluation: 500 owners/month')
add_bullet_ltr(doc, 'Intent: 50 meetings/month')
add_bullet_ltr(doc, 'Purchase: 2-4 contracts/month')

doc.add_page_break()

# 14. Operations and Quality Plan
add_heading_ltr(doc, '14. Operations and Quality Plan', 1)
add_paragraph_ltr(doc, 'The company follows strict quality standards and Saudi Building Code compliance. Each project goes through six phases with documented inspections and approvals.')

add_heading_ltr(doc, 'Quality Standards', 2)
add_bullet_ltr(doc, 'Full compliance with Saudi Building Code (SBC).')
add_bullet_ltr(doc, 'Quality inspections at every finishing phase.')
add_bullet_ltr(doc, 'Documentation of all materials and compliance certificates.')
add_bullet_ltr(doc, 'Final handover with complete technical and legal file.')

add_heading_ltr(doc, 'Insurance Coverage', 2)
add_bullet_ltr(doc, 'Contractors All Risk (CAR) insurance.')
add_bullet_ltr(doc, 'General public liability insurance.')
add_bullet_ltr(doc, 'Hidden defects insurance.')
add_bullet_ltr(doc, 'Professional liability insurance for engineers.')

doc.add_page_break()

# 15. Conclusion
add_heading_ltr(doc, '15. Conclusion and Recommendations', 1)
add_paragraph_ltr(doc, 'The Real Estate Asset Revitalization Company presents an attractive investment opportunity in the Saudi real estate market. The hybrid partnership model eliminates land purchase requirements and reduces initial capital needs while offering attractive returns.')

add_heading_ltr(doc, 'Key Recommendations', 2)
add_bullet_ltr(doc, 'Start with a pilot phase of SAR 25 million and 3-5 benchmark projects.')
add_bullet_ltr(doc, 'Focus on one city (Riyadh or Jeddah) in the first years.')
add_bullet_ltr(doc, 'Build a database of incomplete assets before official launch.')
add_bullet_ltr(doc, 'Partner with certified contractors providing bank guarantees.')
add_bullet_ltr(doc, 'Include binding arbitration clauses in all contracts.')
add_bullet_ltr(doc, 'Develop an ERP system for portfolio and cash flow management.')
add_bullet_ltr(doc, 'Prepare standardized Muhassah contracts protecting all parties.')

add_heading_ltr(doc, 'Investment Decision Matrix', 2)
add_paragraph_ltr(doc, 'This project is suitable for investors who understand that real estate requires patience over 5-7 years, accept a foundational first year, and appreciate calculated risks with 18.5% expected IRR and 2.5x-4x capital return at exit.')

doc.add_page_break()

# 16. Market Enrichments
add_heading_ltr(doc, '16. Market Enrichments', 1)
add_paragraph_ltr(doc, 'The company can explore additional value-adding opportunities:')

add_bullet_ltr(doc, 'Real Estate Sukuk: Asset-backed Islamic bonds for portfolio expansion after building a track record.')
add_bullet_ltr(doc, 'Real Estate Crowdfunding: Partnering with licensed platforms to diversify funding sources.')
add_bullet_ltr(doc, 'Smart Building Technology: KNX and green building solutions to differentiate projects and accelerate sales.')
add_bullet_ltr(doc, 'Legal Execution Procedures: Clear understanding of enforcement procedures through Ministry of Justice platforms.')

doc.add_page_break()

# 17. References
add_heading_ltr(doc, '17. References and Sources', 1)
add_heading_ltr(doc, 'Official Saudi Sources', 2)
add_bullet_ltr(doc, 'Saudi Central Bank (SAMA) - Real estate financing statistics')
add_bullet_ltr(doc, 'Real Estate General Authority (REGA) - Price indices and licenses')
add_bullet_ltr(doc, 'Ministry of Justice - Real estate transaction data')
add_bullet_ltr(doc, 'Raghdan Platform - District-level price data')
add_bullet_ltr(doc, 'Balady Platform - Building permits and violations')

add_heading_ltr(doc, 'Research Reports', 2)
add_bullet_ltr(doc, 'Knight Frank - Saudi Real Estate Report 2026')
add_bullet_ltr(doc, 'KAPSARC - Housing market analysis')
add_bullet_ltr(doc, 'Sands of Wealth - Jeddah Housing Prices 2026')

doc.add_page_break()

# Appendix A
add_heading_ltr(doc, 'Appendix A: Independent Strategic Analysis', 1)
add_paragraph_ltr(doc, 'This appendix provides an independent strategic review of the feasibility study, intended to enhance credibility and demonstrate that the study can withstand external scrutiny.')

add_heading_ltr(doc, 'A.1 Capital Structure and Liquidity', 2)
add_paragraph_ltr(doc, 'The study wisely adopts a flexible capital ceiling of SAR 226 million, with actual operations starting at only SAR 25 million. This phased drawdown approach avoids early capital freezing and reduces investor risk.')

add_heading_ltr(doc, 'A.2 Proof of Concept Strategy', 2)
add_paragraph_ltr(doc, 'Starting with 3-5 benchmark projects over 12-18 months is a prudent approach. Strict milestones should be set before moving to the expansion phase, including achieving 80% sales and SAR 1.5 million net profit.')

add_heading_ltr(doc, 'A.3 Long-Term Value Creation', 2)
add_paragraph_ltr(doc, 'The investment should be viewed as value investing. The true return comes from accumulated asset value, brand recognition, and recurring revenue streams rather than immediate annual distributions.')

doc.add_page_break()

# Appendix B
add_heading_ltr(doc, 'Appendix B: Investor Evaluation', 1)
add_paragraph_ltr(doc, 'This appendix provides a practical investor perspective on the study, highlighting both strengths and concerns.')

add_heading_ltr(doc, 'B.1 Key Value Driver', 2)
add_paragraph_ltr(doc, 'The real value is not in finishing or smart technology, but in the ability to access incomplete buildings before competitors. Strong relationships with owners, appraisers, lawyers, brokers, and banks create a durable competitive advantage.')

add_heading_ltr(doc, 'B.2 Main Concern', 2)
add_paragraph_ltr(doc, 'The study assumes an aggressive project ramp-up: 2 projects in Year 1, 4 in Year 2, 6 in Year 3, and 10 in Year 5. The most challenging part is not execution but acquiring suitable buildings.')

add_heading_ltr(doc, 'B.3 Recommendation', 2)
add_bullet_ltr(doc, 'Keep the administrative team lean - 15 employees at SAR 3.06 million annually is reasonable.')
add_bullet_ltr(doc, 'Position the company as a real estate asset revitalization company, not a smart building technology company.')
add_bullet_ltr(doc, 'Focus geographically in the beginning - start with Jeddah only.')
add_bullet_ltr(doc, 'First project: a 6-12 apartment residential building.')

# Save
doc.save('C:/Users/vip/bonds-global-web/docs/english-study/feasibility-study-english.docx')
print('English feasibility study created successfully')
