from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_rtl(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.right_to_left = True

def add_heading_rtl(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    set_rtl(heading)
    for run in heading.runs:
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        if level == 1:
            run.font.color.rgb = RGBColor(0, 51, 102)
        elif level == 2:
            run.font.color.rgb = RGBColor(212, 168, 83)
    return heading

def add_paragraph_rtl(doc, text, bold=False, italic=False, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    set_rtl(p)
    return p

def add_bullet_rtl(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(11)
    set_rtl(p)
    return p

def add_table_rtl(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            set_rtl(paragraph)
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                run.font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = __import__('docx.oxml', fromlist=['OxmlElement']).OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '0A0F1A')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                set_rtl(paragraph)
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    return table

# Create document
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

# Cover Page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('دراسة جدوى')
run.font.size = Pt(44)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('شركة إحياء الأصول غير المكتملة العقارية')
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(212, 168, 83)
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('شركة مساهمة مقفلة')
run.font.size = Pt(18)
run.font.bold = True
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('المملكة العربية السعودية')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(102, 102, 102)
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('إعداد: Bonds Financial Management & Consulting')
run.font.size = Pt(12)
run.font.italic = True
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('يونيو 2026')
run.font.size = Pt(12)
run.font.italic = True
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('المستشار المالي')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(212, 168, 83)
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('د. طلال بن حسن الزهراني')
run.font.size = Pt(18)
run.font.bold = True
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('+966 56 756 6616')
run.font.size = Pt(14)
run.font.name = 'Arial'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

doc.add_page_break()

# Table of Contents
add_heading_rtl(doc, 'جدول المحتويات', 1)
items = [
    'مقدمة: مناخ الاستثمار والقوة الشرائية',
    '1. الملخص التنفيذي',
    '2. نظرة عامة على المشروع',
    '3. تحليل السوق',
    '4. النموذج التجاري',
    '5. نموذج الأعمال المتكامل',
    '6. الهيكل التنظيمي والفريق',
    '7. الخطة التشغيلية',
    '8. الجداول المالية',
    '9. الجداول المالية التفصيلية',
    '10. تحليل التعادل والحساسية',
    '11. تحليل المخاطر والتحوطات',
    '12. السيناريوهات الثلاثة',
    '13. خطة التسويق والمبيعات',
    '14. خطة التشغيل والجودة',
    '15. الخاتمة والتوصيات',
    '16. إثراءات السوق',
    '17. المراجع والمصادر',
    'ملحق أ: تحليل استراتيجي مستقل',
    'ملحق ب: تقييم المستثمر',
]
for item in items:
    add_bullet_rtl(doc, item)

doc.add_page_break()

# Introduction
add_heading_rtl(doc, 'مقدمة: مناخ الاستثمار والقوة الشرائية', 1)
add_paragraph_rtl(doc, 'في ظل رؤية 2030، يشهد قطاع الاستثمار والتطوير العقاري في المملكة العربية السعودية تحولاً جذرياً. يعد القطاع من أكثر القطاعات حيوية ونضجاً في المنطقة، مدفوعاً بديناميكيات طلب حقيقية بدلاً من المضاربات العشوائية.')
add_paragraph_rtl(doc, 'أُعدت هذه الدراسة بعناية فائقة بناءً على بيانات رسمية من البنك المركزي السعودي (ساما)، الهيئة العامة للعقار (REGA)، وزارة العدل، وتحليلات من Knight Frank و KAPSARC. الهدف هو تقديم رؤية واقعية للسوق تمكن المستثمرين ومالكي المباني من اتخاذ قرارات مستنيرة.')

add_heading_rtl(doc, 'المؤشرات الرئيسية', 2)
rows = [
    ['العائد الإيجاري في الرياض', '8.89%'],
    ['العائد الإيجاري في جدة', '7.89%'],
    ['نمو الصفقات العقارية (2025)', '+14.2%'],
    ['حجم التمويل العقاري', '730 مليار ريال'],
    ['عائد رأس المال عند التخارج', '2.5x - 4x'],
]
add_table_rtl(doc, ['المؤشر', 'القيمة'], rows)

add_paragraph_rtl(doc, 'بينما يمثل العائد الإيجاري دخلاً تشغيلياً سنوياً، فإن العائد الأساسي للمستثمر يأتي من مضاعفة رأس المال عند البيع الاستراتيجي بعد إحياء الأصول. على أفق 10 سنوات، يبلغ IRR المتوقع حوالي 18.5%، مع عائد رأسمالي 2.5x - 4x عند التخارج في السنوات 5-7.')

add_heading_rtl(doc, 'محركات النمو', 2)
add_bullet_rtl(doc, 'المشاريع العملاقة مثل نيوم، البحر الأحمر، الدرعية، وروشن.')
add_bullet_rtl(doc, 'برامج التمويل السكني "سكني" للوحدات المتوسطة (70-120 م²).')
add_bullet_rtl(doc, 'نمو الرواتب والهجرة نحو المدن الكبرى.')
add_bullet_rtl(doc, 'العوائد الإيجارية المستقرة والجاذبة.')

doc.add_page_break()

# 1. Executive Summary
add_heading_rtl(doc, '1. الملخص التنفيذي', 1)
add_heading_rtl(doc, '1.1 الفكرة الأساسية', 2)
add_paragraph_rtl(doc, 'تأسست شركة إحياء الأصول غير المكتملة العقارية على رؤية متمثلة في إحياء المباني غير المكتملة دون شراء أراضٍ. يقوم النموذج على عقود شراكة (محاصة) مع مالكي المباني، تتحمل فيها الشركة تكاليف التشطيب والتسويق والبيع، ثم تقاسم الأرباح مع المالك الأصلي.')

add_heading_rtl(doc, 'نقاط القوة الرئيسية', 2)
add_bullet_rtl(doc, 'لا يتطلب شراء أراضٍ - مما يقلل مخاطر الدخول للسوق ورأس المال المطلوب.')
add_bullet_rtl(doc, 'يحتوي السوق السعودي على آلاف المباني غير المكتملة.')
add_bullet_rtl(doc, 'هوامش ربح صحية تتراوح بين 8-12% مقارنة بالتطوير العقاري التقليدي.')
add_bullet_rtl(doc, 'السنة الأولى تأسيسية؛ تبدأ الإيرادات الحقيقية من السنة الثانية.')
add_bullet_rtl(doc, 'رأس المال المطلوب فعلياً 25 مليون ريال فقط، مع سقف مرن يصل إلى 226 مليون ريال.')
add_bullet_rtl(doc, 'القيمة الحقيقية تكمن في تراكم محفظة الأصول والعلامة التجارية والإيرادات المتكررة.')
add_bullet_rtl(doc, 'العائد الأساسي يأتي من التخارج الاستراتيجي ورفع القيمة السوقية للشركة.')

add_heading_rtl(doc, 'المؤشرات المالية الرئيسية', 2)
rows = [
    ['رأس المال المصرح', '200,000,000 ريال'],
    ['عدد الأسهم', '40,000 سهم'],
    ['القيمة الاسمية للسهم', '5,000 ريال'],
    ['القيمة الفعلية للسهم', '5,650 ريال'],
    ['إجمالي الاستثمار المستهدف', '226,000,000 ريال'],
    ['IRR المتوقع (10 سنوات)', '18.5% - 20%'],
    ['عائد رأس المال عند الخروج', '2.5x - 4x'],
    ['فترة استرداد رأس المال', '6 سنوات'],
]
add_table_rtl(doc, ['المؤشر', 'القيمة'], rows)

doc.add_page_break()

# 2. Project Overview
add_heading_rtl(doc, '2. نظرة عامة على المشروع', 1)
add_heading_rtl(doc, '2.1 الشكل القانوني', 2)
add_paragraph_rtl(doc, 'شركة مساهمة مقفلة مرخصة من وزارة التجارة وهيئة العقار (REGA)، وفقاً لنظام الشركات السعودي. يتيح هذا الشكل جمع رأس المال من المساهمين المؤسسين مع الحفاظ على مرونة الإدارة وسرعة اتخاذ القرار.')

add_heading_rtl(doc, '2.2 هيكل رأس المال', 2)
rows = [
    ['البند', 'المبلغ (ريال)'],
    ['رأس المال المصرح', '200,000,000'],
    ['عدد الأسهم', '40,000'],
    ['القيمة الاسمية للسهم', '5,000'],
    ['القيمة الفعلية للسهم', '5,650'],
    ['علاوة الإصدار ومصاريف التأسيس', '26,000,000'],
    ['إجمالي الاستثمار المستهدف', '226,000,000'],
]
add_table_rtl(doc, ['البند', 'المبلغ (ريال)'], rows)

add_heading_rtl(doc, '2.3 استخدامات رأس المال', 2)
rows = [
    ['الاستخدام', 'المبلغ (مليون ريال)', 'النسبة'],
    ['استثمارات في مشاريع (شراء + تشطيب)', '150', '66%'],
    ['رأس المال العامل', '40', '18%'],
    ['تكنولوجيا وأنظمة', '10', '4%'],
    ['تأسيس وتسويق وعلاقات', '15', '7%'],
    ['احتياطي سيولة طوارئ', '11', '5%'],
]
add_table_rtl(doc, ['الاستخدام', 'المبلغ', 'النسبة'], rows)

add_heading_rtl(doc, '2.4 آلية سحب رأس المال', 2)
add_bullet_rtl(doc, 'الدفعة الأولى (30%): التأسيس وشراء 3-5 مبانٍ أولى.')
add_bullet_rtl(doc, 'الدفعة الثانية (30%): إنجاز أول مشروعين وتحقيق أولى الإيرادات.')
add_bullet_rtl(doc, 'الدفعة الثالثة (25%): الوصول إلى 6 مشاريع نشطة وتحقيق التعادل التشغيلي.')
add_bullet_rtl(doc, 'الدفعة الرابعة (15%): احتياطي مرن يُطلق حسب الحاجة أو يُعاد للمساهمين.')

doc.add_page_break()

# 3. Market Analysis
add_heading_rtl(doc, '3. تحليل السوق', 1)
add_paragraph_rtl(doc, 'يشهد السوق العقاري السعودي تحولاً هيكلياً عميقاً. المدن الكبرى مثل الرياض وجدة تسجل عوائد إيجارية استثنائية مقارنة بالعواصم العالمية. الطلب على الوحدات الجاهزة يفوق العرض، مما يخلق فرصة كبيرة لإحياء الأصول غير المكتملة.')

add_heading_rtl(doc, '3.1 المناطق المستهدفة', 2)
add_paragraph_rtl(doc, 'الرياض:', bold=True)
add_bullet_rtl(doc, 'الشمال الفاخر: العقيق، حطين، الملقا - عوائد إيجارية مرتفعة وطلب تنفيذي.')
add_bullet_rtl(doc, 'الشمال الشرقي: المونسية، الرمال، الياسمين، النرجس - كثافة سكانية عالية وبنية تحتية حديثة.')

add_paragraph_rtl(doc, 'جدة:', bold=True)
add_bullet_rtl(doc, 'الساحل الشمالي: أبحر الشمالية، الصواري - إمكانية سكنية وسياحية طويلة المدى.')
add_bullet_rtl(doc, 'الوسط الحيوي: الحمراء، النعيم، الزهراء، السلامة - عوائد إيجارية ثابتة.')

add_heading_rtl(doc, '3.2 الميزة التنافسية', 2)
add_paragraph_rtl(doc, 'المطورون التقليديون يتطلب نموذجهم شراء أراضٍ ورأس مال ضخم. ميزتنا تكمن في عدم الحاجة لشراء الأرض، والتركيز على الأصول غير المكتملة التي لا تجذب المطورين الكبار، مما يتيح سرعة تنفيذ وتكاليف أقل.')

doc.add_page_break()

# 4. Commercial Model
add_heading_rtl(doc, '4. النموذج التجاري', 1)
add_paragraph_rtl(doc, 'تعمل الشركة بنموذج شراكة محاصة متوازن يجمع بين مصادر متعددة للإيرادات:')

add_heading_rtl(doc, 'نماذج الشراكة', 2)
rows = [
    ['النموذج', 'الوصف'],
    ['المحاصة الكاملة', 'تتولى الشركة التشطيب والتسويق مقابل نسبة من الأرباح'],
    ['الشراء الجزئي', 'تشتري الشركة جزءاً من الوحدات بعد التشطيب'],
    ['رسوم الإدارة', 'رسوم ثابتة مقابل التشطيب والتسويق'],
    ['النموذج المختلط', 'مزيج من المحاصة والشراء ورسوم الإدارة'],
]
add_table_rtl(doc, ['النموذج', 'الوصف'], rows)

add_heading_rtl(doc, 'مصادر الإيرادات', 2)
rows = [
    ['مصدر الإيراد', 'النسبة المتوقعة'],
    ['إيرادات المحاصة', '60%'],
    ['أرباح إعادة البيع', '25%'],
    ['رسوم الإدارة', '10%'],
    ['خدمات استشارية', '5%'],
]
add_table_rtl(doc, ['مصدر الإيراد', 'النسبة'], rows)

doc.add_page_break()

# 5. Integrated Business Model
add_heading_rtl(doc, '5. النموذج المتكامل للأعمال', 1)
add_paragraph_rtl(doc, 'يجمع النموذج المتكامل بين مصادر الإيرادات المتعددة وكفاءة التكاليف لخلق عملية مستدامة وقابلة للتوسع. من خلال الاستفادة من التقنية والشبكات المهنية والعمليات الموحدة، يمكن للشركة تكرار النجاح عبر مشاريع متعددة.')

add_bullet_rtl(doc, 'تقييم مركزي للمشاريع ومراجعة قانونية.')
add_bullet_rtl(doc, 'مواصفات تشطيب موحدة وإدارة مقاولين منظمة.')
add_bullet_rtl(doc, 'منصة تسويق ومبيعات موحدة.')
add_bullet_rtl(doc, 'اختيار الأصول وتسعيرها بناءً على البيانات.')

doc.add_page_break()

# 6. Organizational Structure
add_heading_rtl(doc, '6. الهيكل التنظيمي والفريق', 1)
add_paragraph_rtl(doc, 'تم تصميم الهيكل الإداري ليكون مرناً وفعالاً، مع بدء تشغيلي يعتمد على 15 موظفاً في السنوات الأولى.')

add_heading_rtl(doc, 'الأدوار الرئيسية', 2)
rows = [
    ['الدور', 'المسؤوليات'],
    ['الرئيس التنفيذي', 'الاستراتيجية، علاقات المستثمرين، القرارات الكبرى'],
    ['المدير التشغيلي', 'إدارة المشاريع، التنسيق مع المقاولين، ضمان الجودة'],
    ['المدير المالي', 'التدفقات النقدية، التقارير المالية، التخطيط المالي'],
    ['المدير التجاري', 'التسويق والمبيعات، علاقات المالكين، تطوير الأعمال'],
    ['المدير القانوني', 'العقود، الامتثال، حل النزاعات'],
    ['مدير الموارد البشرية والإدارة', 'الموظفين والإجراءات الإدارية'],
]
add_table_rtl(doc, ['الدور', 'المسؤوليات'], rows)

add_heading_rtl(doc, 'خطة التوظيف', 2)
rows = [
    ['السنة', 'عدد الموظفين', 'الملاحظات'],
    ['1', '10 - 15', 'فريق أساسي للمرحلة الأولى'],
    ['2', '15', 'تثبيت الهيكل التشغيلي'],
    ['3', '18 - 20', 'التوسع في عدد المشاريع'],
    ['5', '25 - 30', 'دعم محفظة 10 مشاريع سنوياً'],
]
add_table_rtl(doc, ['السنة', 'الموظفين', 'الملاحظات'], rows)

doc.add_page_break()

# 7. Operational Plan
add_heading_rtl(doc, '7. الخطة التشغيلية', 1)
add_heading_rtl(doc, '7.1 مراحل المشروع الواحد', 2)
rows = [
    ['المرحلة', 'المدة', 'الأنشطة'],
    ['التقييم', '4 أسابيع', 'فحص هيكلي، مراجعة قانونية، تقدير تكلفة'],
    ['التعاقد', '2 - 4 أسابيع', 'عقد محاصة، استصدار الموافقات'],
    ['التصميم', '4 - 8 أسابيع', 'التصاميم، التراخيص البلدية'],
    ['التشطيب', '16 - 24 أسبوع', 'اختيار مقاول، إشراف يومي'],
    ['التسويق والبيع', '12 - 24 أسبوع', 'عرض الوحدات، عقود البيع'],
    ['التسليم والتسوية', '4 - 8 أسابيع', 'تسليم الوحدات، نقل الملكية، توزيع الأرباح'],
]
add_table_rtl(doc, ['المرحلة', 'المدة', 'الأنشطة'], rows)

add_heading_rtl(doc, '7.2 مؤشرات الأداء الرئيسية', 2)
rows = [
    ['المؤشر', 'الهدف', 'التحذير', 'الخطر'],
    ['نسبة إنجاز التشطيب', '>80% في الشهر الرابع', '<60%', 'تأخر المقاول'],
    ['معدل بيع الوحدات', '>30% في الشهر الثاني', '<10%', 'ضعف الطلب'],
    ['تكلفة التشطيب/المتوقع', '<105%', '>115%', 'عيوب مخفية'],
    ['رضا العملاء (NPS)', '>50', '<30', 'مشاكل جودة'],
    ['مدة البيع', '<6 أشهر', '>9 أشهر', 'ركود السوق'],
]
add_table_rtl(doc, ['المؤشر', 'الهدف', 'التحذير', 'الخطر'], rows)

doc.add_page_break()

# 8. Financial Tables
add_heading_rtl(doc, '8. الجداول المالية', 1)
add_heading_rtl(doc, '8.1 افتراضات النموذج المالي', 2)
rows = [
    ['الافتراض', 'القيمة'],
    ['متوسط مساحة المبنى', '1,500 م²'],
    ['عدد الوحدات في المبنى', '6 - 8 وحدات'],
    ['تكلفة التشطيب', '1,200 - 1,800 ريال/م²'],
    ['سعر البيع', '2,500 - 3,500 ريال/م²'],
    ['نسبة المحاصة للشركة', '50%'],
]
add_table_rtl(doc, ['الافتراض', 'القيمة'], rows)

add_heading_rtl(doc, '8.2 التكلفة التقديرية للمشروع الواحد', 2)
rows = [
    ['البند', 'التكلفة (ريال)', 'النسبة'],
    ['تكاليف التشطيب', '2,250,000', '55%'],
    ['التراخيص والرسوم', '150,000', '4%'],
    ['الاستشارات الهندسية', '200,000', '5%'],
    ['التسويق والمبيعات', '300,000', '7%'],
    ['احتياطي المخاطر (15%)', '450,000', '11%'],
    ['التكاليف التشغيلية', '250,000', '6%'],
    ['التكاليف المالية', '200,000', '5%'],
    ['رسوم البيع والنقل', '375,000', '9%'],
    ['الإجمالي', '4,175,000', '100%'],
]
add_table_rtl(doc, ['البند', 'التكلفة', 'النسبة'], rows)

add_heading_rtl(doc, '8.3 قائمة الدخل لـ 5 سنوات (النموذج المختلط)', 2)
rows = [
    ['البند', 'السنة 1', 'السنة 2', 'السنة 3', 'السنة 4', 'السنة 5'],
    ['إجمالي الإيرادات (مليون ريال)', '3.2', '6.4', '9.6', '12.8', '16.0'],
    ['التكاليف المباشرة', '-0.4', '-0.8', '-1.2', '-1.6', '-2.0'],
    ['تكاليف التشغيل', '-3.2', '-3.2', '-3.5', '-3.8', '-4.0'],
    ['صافي الربح قبل الزكاة', '-0.5', '2.4', '5.0', '7.7', '10.4'],
]
add_table_rtl(doc, ['البند', 'Y1', 'Y2', 'Y3', 'Y4', 'Y5'], rows)

doc.add_page_break()

# 9. Detailed Financial Tables
add_heading_rtl(doc, '9. الجداول المالية التفصيلية', 1)
add_paragraph_rtl(doc, 'يتضمن النموذج المالي التفصيلي التدفقات النقدية الشهرية للسنة الأولى، وتفاصيل تكاليف المشروع، ومؤشرات العائد. يفترض النموذج بدءاً تدريجياً بمشروعين إلى ثلاثة في السنة الأولى، والوصول إلى 10 مشاريع في السنة الخامسة.')

add_heading_rtl(doc, '9.1 مؤشرات العائد', 2)
rows = [
    ['المؤشر', 'القيمة'],
    ['عائد رأس المال عند الخروج (سنوات 5-7)', '2.5x - 4x'],
    ['IRR المتوقع (10 سنوات)', '18.5%+'],
    ['القيمة السوقية المتوقعة (السنة 5)', '75 - 94 مليون ريال'],
    ['فترة استرداد رأس المال', '6 سنوات'],
]
add_table_rtl(doc, ['المؤشر', 'القيمة'], rows)

add_heading_rtl(doc, '9.2 الإسقاط طويل المدى (سنوات 6-10)', 2)
rows = [
    ['السنة', 'المشاريع', 'الإيرادات (مليون ريال)', 'صافي الربح (مليون ريال)'],
    ['6', '12', '19.2', '6.0+'],
    ['7', '14', '22.4', '8.0+'],
    ['8', '16', '25.6', '10.0+'],
    ['9', '18', '28.8', '12.0+'],
    ['10', '20', '32.0', '14.0+'],
]
add_table_rtl(doc, ['السنة', 'المشاريع', 'الإيرادات', 'صافي الربح'], rows)

doc.add_page_break()

# 10. Break-even and Sensitivity
add_heading_rtl(doc, '10. تحليل التعادل والحساسية', 1)
add_heading_rtl(doc, '10.1 نقطة التعادل', 2)
add_paragraph_rtl(doc, 'في النموذج المختلط، تصل الشركة للتعادل التشغيلي بحوالي 5-6 مشاريع سنوياً. هذا يفترض مزيجاً متوازناً من إيرادات المحاصة والشراء الجزئي ورسوم الإدارة.')

add_heading_rtl(doc, '10.2 تحليل الحساسية', 2)
rows = [
    ['تغير السعر', 'صافي الربح السنة 3', 'صافي الربح السنة 5'],
    ['-20%', '1.5 مليون', '5.5 مليون'],
    ['-10%', '3.2 مليون', '7.8 مليون'],
    ['الأساس (0%)', '5.0 مليون', '10.4 مليون'],
    ['+10%', '6.8 مليون', '13.0 مليون'],
    ['+20%', '8.5 مليون', '15.5 مليون'],
]
add_table_rtl(doc, ['تغير السعر', 'Y3', 'Y5'], rows)

add_paragraph_rtl(doc, 'يظهر النموذج مرونة حتى مع انخفاض 20% في الأسعار أو ارتفاع 30% في التكاليف، حيث يبقى مربحاً من السنة الثالثة فصاعداً.')

doc.add_page_break()

# 11. Risk Analysis
add_heading_rtl(doc, '11. تحليل المخاطر والتحوطات', 1)
rows = [
    ['المخاطر', 'الاحتمالية', 'التأثير', 'التحوط'],
    ['صعوبة الحصول على أصول جيدة', 'عالية', 'توقف النمو', 'شبكة علاقات وقاعدة بيانات ومصفين'],
    ['عيوب إنشائية مخفية', 'عالية', 'خسارة مالية كبيرة', 'فحص هندسي + تأمين + احتياطي 15%'],
    ['تأخر بيع الوحدات', 'متوسطة', 'حمل تكاليف', 'تسعير تنافسي + عقود مسبقة البيع'],
    ['تأخر المقاول', 'متوسطة', 'تكاليف إضافية', 'كفالة بنكية + عقوبات تأخير'],
    ['تغير أسعار مواد البناء', 'متوسطة', 'ارتفاع التكاليف', 'عقود ثابتة السعر'],
    ['نزاعات قانونية مع المالك', 'متوسطة', 'توقف المشروع', 'عقود واضحة + تحكيم ملزم'],
    ['تراجع أسعار العقار', 'منخفضة-متوسطة', 'انخفاض الهوامش', 'تنويع جغرافي وفئات استهداف'],
]
add_table_rtl(doc, ['المخاطر', 'الاحتمالية', 'التأثير', 'التحوط'], rows)

add_heading_rtl(doc, 'الاحتياطي المقترح للمخاطر', 2)
rows = [
    ['نوع الاحتياطي', 'المبلغ'],
    ['احتياطي عيوب إنشائية', '225,000 ريال/مشروع'],
    ['احتياطي تأخر البيع', '150,000 ريال/مشروع'],
    ['احتياطي عام', '2,500,000 ريال'],
    ['الإجمالي المقترح', '3,000,000 ريال'],
]
add_table_rtl(doc, ['الاحتياطي', 'المبلغ'], rows)

doc.add_page_break()

# 12. Three Scenarios
add_heading_rtl(doc, '12. السيناريوهات الثلاثة', 1)
rows = [
    ['السيناريو', 'مشاريع السنة 1', 'تكلفة التشطيب', 'مدة البيع', 'صافي الربح السنة 5'],
    ['متشائم', '1', '+20% عن المتوقع', '+12 شهراً', '3.0 مليون ريال'],
    ['متوقع', '2 - 3', 'كما هو متوقع', '3 - 6 أشهر', '10.4 مليون ريال'],
    ['متفائل', '4', '-10% عن المتوقع', '3 أشهر', '14.0 مليون ريال'],
]
add_table_rtl(doc, ['السيناريو', 'Y1', 'التكلفة', 'البيع', 'Y5 Profit'], rows)

add_heading_rtl(doc, 'التعادل حسب السيناريو', 2)
rows = [
    ['السيناريو', 'التعادل السنوي', 'التعادل التراكمي'],
    ['متشائم', '18 مشروع', '54 مشروع'],
    ['متوقع', '6 مشاريع', '20 مشروع'],
    ['متفائل', '10 مشاريع', '20 مشروع'],
]
add_table_rtl(doc, ['السيناريو', 'سنوي', 'تراكمي'], rows)

doc.add_page_break()

# 13. Marketing and Sales Plan
add_heading_rtl(doc, '13. خطة التسويق والمبيعات', 1)
add_heading_rtl(doc, '13.1 استراتيجية التسويق الرقمي', 2)
rows = [
    ['القناة', 'الهدف', 'الميزانية الشهرية'],
    ['Google Ads', 'الوصول لمالكي العقارات غير المكتملة', '10,000 ريال'],
    ['Snapchat Ads', 'الفئة العمرية 25-45', '8,000 ريال'],
    ['Instagram / TikTok', 'عرض مشاريع "قبل وبعد"', '5,000 ريال'],
    ['SEO', 'الظهور في "إحياء المباني"', '3,000 ريال'],
    ['LinkedIn', 'الوصول للمستثمرين والمالكين', '2,000 ريال'],
]
add_table_rtl(doc, ['القناة', 'الهدف', 'الميزانية'], rows)

add_heading_rtl(doc, '13.2 قمع المبيعات', 2)
add_bullet_rtl(doc, 'الوعي: 100,000 شخص/شهر')
add_bullet_rtl(doc, 'الاهتمام: 5,000 شخص/شهر')
add_bullet_rtl(doc, 'التقييم: 500 مالك/شهر')
add_bullet_rtl(doc, 'النية: 50 اجتماع/شهر')
add_bullet_rtl(doc, 'الشراء: 2-4 عقد/شهر')

doc.add_page_break()

# 14. Operations and Quality Plan
add_heading_rtl(doc, '14. خطة التشغيل والجودة', 1)
add_paragraph_rtl(doc, 'تلتزم الشركة بمعايير جودة صارمة والكود السعودي للبناء. يمر كل مشروع بست مراحل مع تفتيشات وموافقات موثقة.')

add_heading_rtl(doc, 'معايير الجودة', 2)
add_bullet_rtl(doc, 'الالتزام الكامل بكود البناء السعودي (SBC).')
add_bullet_rtl(doc, 'فحوصات جودة دورية في كل مرحلة من مراحل التشطيب.')
add_bullet_rtl(doc, 'توثيق جميع المواد المستخدمة وشهادات المطابقة.')
add_bullet_rtl(doc, 'التسليم بملف فني وقانوني كامل.')

add_heading_rtl(doc, 'التغطية التأمينية', 2)
add_bullet_rtl(doc, 'تأمين المقاولين الشامل (CAR).')
add_bullet_rtl(doc, 'تأمين المسؤولية المدنية العامة.')
add_bullet_rtl(doc, 'تأمين العيوب الخفية.')
add_bullet_rtl(doc, 'تأمين المسؤولية المهنية للمهندسين.')

doc.add_page_break()

# 15. Conclusion
add_heading_rtl(doc, '15. الخاتمة والتوصيات', 1)
add_paragraph_rtl(doc, 'تقدم شركة إحياء الأصول غير المكتملة العقارية فرصة استثمارية فريدة في السوق العقاري السعودي. يعتمد النموذج التجاري على الشراكة المحاصة مع مالكي العقارات غير المكتملة، مما يلغي الحاجة إلى شراء الأراضي ويقلل رأس المال المطلوب.')

add_heading_rtl(doc, 'التوصيات الرئيسية', 2)
add_bullet_rtl(doc, 'البدء بمرحلة تجريبية بـ 25 مليون ريال و3-5 مشاريع قياسية.')
add_bullet_rtl(doc, 'التركيز على مدينة واحدة (الرياض أو جدة) في السنوات الأولى.')
add_bullet_rtl(doc, 'بناء قاعدة بيانات للأصول غير المكتملة قبل الإطلاق الرسمي.')
add_bullet_rtl(doc, 'إبرام شراكات مع مقاولين معتمدين يقدمون كفالات بنكية.')
add_bullet_rtl(doc, 'إدراج شرط تحكيم ملزم في جميع عقود المحاصة.')
add_bullet_rtl(doc, 'تطوير نظام ERP لإدارة المحفظة والتدفقات النقدية.')
add_bullet_rtl(doc, 'إعداد عقد محاصة موحد يحفظ حقوق الشركة والمالك.')

add_heading_rtl(doc, 'مصفوفة قرار الاستثمار', 2)
add_paragraph_rtl(doc, 'هذا المشروع مناسب للمستثمرين الذين يفهمون أن العقارات تحتاج صبراً 5-7 سنوات، ويقبلون سنة أولى تأسيسية، ويقدرون المخاطر المحسوبة مع IRR متوقع 18.5% وعائد رأسمالي 2.5x-4x عند الخروج.')

doc.add_page_break()

# 16. Market Enrichments
add_heading_rtl(doc, '16. إثراءات السوق', 1)
add_paragraph_rtl(doc, 'يمكن للشركة استكشاف فرص إضافية لخلق قيمة:')

add_bullet_rtl(doc, 'الصكوك العقارية: سندات إسلامية مدعومة بالأصول للتوسع بعد بناء سجل حافل.')
add_bullet_rtl(doc, 'التمويل الجماعي العقاري: شراكات مع منصات مرخصة لتنويع مصادر التمويل.')
add_bullet_rtl(doc, 'تقنية المباني الذكية: حلول KNX والبناء الأخضر للتميز وتسريع المبيعات.')
add_bullet_rtl(doc, 'الإجراءات القضائية: فهم واضح لإجراءات التنفيذ عبر منصات وزارة العدل.')

doc.add_page_break()

# 17. References
add_heading_rtl(doc, '17. المراجع والمصادر', 1)
add_heading_rtl(doc, 'مصادر رسمية سعودية', 2)
add_bullet_rtl(doc, 'البنك المركزي السعودي (ساما) - إحصائيات التمويل العقاري')
add_bullet_rtl(doc, 'الهيئة العامة للعقار (REGA) - مؤشرات الأسعار والتراخيص')
add_bullet_rtl(doc, 'وزارة العدل - بيانات الصفقات العقارية')
add_bullet_rtl(doc, 'منصة رغدان - أسعار المربع السكني')
add_bullet_rtl(doc, 'منصة بلدي - رخص البناء والمخالفات')

add_heading_rtl(doc, 'تقارير بحثية', 2)
add_bullet_rtl(doc, 'Knight Frank - Saudi Real Estate Report 2026')
add_bullet_rtl(doc, 'KAPSARC - تحليل السوق الإسكاني')
add_bullet_rtl(doc, 'Sands of Wealth - Jeddah Housing Prices 2026')

doc.add_page_break()

# Appendix A
add_heading_rtl(doc, 'ملحق أ: تحليل استراتيجي مستقل', 1)
add_paragraph_rtl(doc, 'يقدم هذا الملحق مراجعة استراتيجية مستقلة لدراسة الجدوى، بهدف تعزيز المصداقية وإظهار أن الدراسة قابلة لتحمل التدقيق الخارجي.')

add_heading_rtl(doc, 'أ.1 هيكل رأس المال والسيولة', 2)
add_paragraph_rtl(doc, 'تعتمد الدراسة بحكمة على سقف رأس مال مرن قدره 226 مليون ريال، مع بدء العمليات الفعلية بمبلغ 25 مليون ريال فقط. يجنب هذا النهج تجميد رأس المال في المراحل المبكرة ويقلل من مخاطر المستثمرين.')

add_heading_rtl(doc, 'أ.2 استراتيجية إثبات النموذج', 2)
add_paragraph_rtl(doc, 'البدء بـ 3-5 مشاريع قياسية على مدار 12-18 شهراً هو نهج حكيم. يجب وضع معالم صارمة قبل الانتقال للمرحلة التوسعية، بما في ذلك تحقيق 80% مبيعات وصافي ربح 1.5 مليون ريال.')

add_heading_rtl(doc, 'أ.3 خلق القيمة طويلة المدى', 2)
add_paragraph_rtl(doc, 'يجب النظر إلى هذا الاستثمار كنموذج استثمار قيمة (Value Investing). العائد الحقيقي يأتي من تراكم قيمة الأصول والعلامة التجارية وتدفقات الإيرادات المتكررة وليس من التوزيعات السنوية الفورية.')

doc.add_page_break()

# Appendix B
add_heading_rtl(doc, 'ملحق ب: تقييم المستثمر', 1)
add_paragraph_rtl(doc, 'يقدم هذا الملحق رؤية عملانية من منظور مستثمر عملاني، مع تسليط الضوء على نقاط القوة والقلق.')

add_heading_rtl(doc, 'ب.1 محرك القيمة الرئيسي', 2)
add_paragraph_rtl(doc, 'القيمة الحقيقية ليست في التشطيب أو التقنية الذكية، بل في القدرة على الوصول للمباني غير المكتملة قبل المنافسين. العلاقات القوية مع المالكين، المصفين، المحامين، الوسطاء والبنوك تخلق ميزة تنافسية مستدامة.')

add_heading_rtl(doc, 'ب.2 القلق الرئيسي', 2)
add_paragraph_rtl(doc, 'تفترض الدراسة تسارعاً سريعاً: مشروعين في السنة الأولى، 4 في الثانية، 6 في الثالثة، و10 في الخامسة. الجزء الأكثر تحدياً ليس التنفيذ، بل الحصول على المباني المناسبة.')

add_heading_rtl(doc, 'ب.3 التوصيات', 2)
add_bullet_rtl(doc, 'الإبقاء على الهيكل الإداري مرناً - 15 موظفاً بتكلفة 3.06 مليون ريال سنوياً هو أمر معقول.')
add_bullet_rtl(doc, 'تموضع الشركة كشركة إحياء أصول عقارية، وليست شركة مباني ذكية.')
add_bullet_rtl(doc, 'التركيز جغرافياً في البداية - البدء بجدة فقط.')
add_bullet_rtl(doc, 'المشروع الأول: عمارة سكنية 6-12 شقة.')

# Save
doc.save('C:/Users/vip/bonds-global-web/docs/english-study/feasibility-study-arabic.docx')
print('Arabic feasibility study created successfully')
